from __future__ import annotations

import json
import math
import os
import tempfile
from dataclasses import asdict, is_dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from pfi_os.approvals import StrategyApprovalRegistry
from pfi_os.app.dashboard import backtest_result_judgements, buy_and_hold_metrics
from pfi_os.analysis import bootstrap_equity_robustness, build_strategy_diagnostics, robustness_summary_rows
from pfi_os.backtest import BacktestResult
from pfi_os.config import report_date_dir
from pfi_os.data.validation import CrossSourceValidationResult
from pfi_os.risk import DecisionQualityResult, RiskGateResult, evaluate_decision_quality, evaluate_research_risk_gates
from pfi_os.storage import atomic_write_json
from pfi_os.strategies.profiles import RETURN_SOURCE_TAXONOMY, get_strategy_profile
from pfi_os.strategies import (
    collect_strategy_code_quality_reports,
    collect_strategy_smoke_tests,
    evaluate_strategy_readiness_gate,
    parse_strategy_profile_candidate,
)


def export_backtest_docx(
    result: BacktestResult,
    output_path: Path | str | None = None,
    report_name: str = "BacktestReport",
    data_quality_report=None,
    cross_validation_result: CrossSourceValidationResult | None = None,
) -> Path:
    output = Path(output_path) if output_path is not None else unique_report_path(report_name, suffix=".docx")
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_document(doc)
    _add_title(doc)
    _add_research_notice(doc)

    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph("Executive Summary")
    _add_executive_summary(doc, result)

    data_quality_status = _data_quality_status(data_quality_report)
    cross_validation_status = getattr(cross_validation_result, "status", None) if cross_validation_result is not None else None
    risk_gate_result = evaluate_research_risk_gates(
        metrics=result.metrics,
        data_quality_status=data_quality_status,
    )
    decision_quality = evaluate_decision_quality(
        result=result,
        risk_gate=risk_gate_result,
        data_quality_status=data_quality_status,
        cross_validation_status=cross_validation_status,
    )

    doc.add_heading("决策质量摘要", level=1)
    doc.add_paragraph("Decision Quality Score")
    _add_decision_quality_section(doc, decision_quality)

    report_evidence = report_evidence_summary(
        result,
        data_quality_report=data_quality_report,
        cross_validation_result=cross_validation_result,
        risk_gate_result=risk_gate_result,
        decision_quality=decision_quality,
    )
    doc.add_heading("报告证据层", level=1)
    doc.add_paragraph("Report Evidence Layer")
    _add_report_evidence_section(doc, report_evidence)

    doc.add_heading("策略说明", level=1)
    doc.add_paragraph("Strategy Description")
    _add_strategy_section(doc, result)

    doc.add_heading("研究风险闸门", level=1)
    doc.add_paragraph("Research Risk Gate")
    _add_risk_gate_section(doc, risk_gate_result)

    doc.add_heading("核心指标", level=1)
    doc.add_paragraph("Key Metrics")
    metrics = dict(result.metrics)
    buy_hold_metrics = _buy_and_hold_metrics(result)
    metrics.update(buy_hold_metrics)
    _add_metrics_table(doc, metrics)
    _add_buy_and_hold_explanation(doc, buy_hold_metrics)

    doc.add_heading("策略诊断", level=1)
    doc.add_paragraph("Strategy Diagnostics")
    _add_strategy_diagnostics_section(doc, result)

    doc.add_heading("图表", level=1)
    doc.add_paragraph("Charts")
    doc.add_paragraph("本节展示策略权益曲线、回撤曲线，以及单标的回测中的价格与买卖点。")
    doc.add_paragraph("This section shows the strategy equity curve, drawdown curve, and price with buy/sell markers for single-symbol backtests.")
    _add_charts(doc, result)

    doc.add_heading("Bootstrap 鲁棒性验证", level=1)
    doc.add_paragraph("Bootstrap Robustness Validation")
    _add_bootstrap_robustness_section(doc, result, target_return=buy_hold_metrics.get("buy_hold_total_return", 0.0))

    doc.add_heading("数据质量摘要", level=1)
    doc.add_paragraph("Data Quality Summary")
    _add_data_quality_summary(doc, data_quality_report)

    doc.add_heading("多源交叉校验摘要", level=1)
    doc.add_paragraph("Cross-Source Validation Summary")
    _add_cross_validation_summary(doc, cross_validation_result)

    doc.add_heading("最近交易", level=1)
    doc.add_paragraph("Recent Trades")
    _add_dataframe_table(doc, result.trades.tail(20), max_rows=20)

    doc.add_heading("权益曲线尾部", level=1)
    doc.add_paragraph("Equity Curve Tail")
    _add_dataframe_table(doc, result.equity_curve.tail(20), max_rows=20)

    doc.add_heading("运行配置与追溯信息", level=1)
    doc.add_paragraph("Run Configuration And Traceability")
    _add_metadata_explanation(doc, result)
    doc.add_paragraph("原始元数据 JSON 如下。")
    doc.add_paragraph("Raw metadata JSON is preserved below.")
    doc.add_paragraph(json.dumps(result.metadata, ensure_ascii=False, indent=2, default=str))

    _save_docx_atomic(doc, output)
    metadata_path = output.with_name(output.stem.replace(report_name, "RunMetadata", 1).replace("BacktestReport", "RunMetadata", 1) + ".json")
    atomic_write_json(
        metadata_path,
        {
            "metrics": result.metrics,
            "metadata": result.metadata,
            "risk_gate": asdict(risk_gate_result),
            "decision_quality": asdict(decision_quality),
            "report_evidence": report_evidence,
        },
        default=str,
    )
    return output


def export_backtest_html(result: BacktestResult, output_path: Path | str | None = None) -> Path:
    """Compatibility wrapper that now exports Word reports."""
    return export_backtest_docx(result, output_path=output_path)


def export_experiment_docx(
    summary_path: Path | str,
    output_path: Path | str | None = None,
    report_name: str = "ExperimentResearchReport",
) -> Path:
    from pfi_os.reports.catalog import load_experiment_detail

    detail = load_experiment_detail(summary_path)
    output = Path(output_path) if output_path is not None else unique_report_path(report_name, suffix=".docx")
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_document(doc)
    _add_experiment_title(doc, str(detail["experiment"]))
    _add_research_notice(doc)

    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph("Executive Summary")
    _add_experiment_executive_summary(doc, detail)

    risk_gate = _risk_gate_from_payload(detail.get("risk_gate", {}))
    decision_quality = evaluate_decision_quality(
        metrics=detail.get("best_run", {}),
        risk_gate=risk_gate,
        stability=detail.get("stability"),
        train_test=detail.get("train_test_validation"),
        walk_forward=detail.get("walk_forward_validation"),
    )
    doc.add_heading("决策质量摘要", level=1)
    doc.add_paragraph("Decision Quality Score")
    _add_decision_quality_section(doc, decision_quality)

    doc.add_heading("研究风险闸门", level=1)
    doc.add_paragraph("Research Risk Gate")
    _add_risk_gate_section(doc, risk_gate)

    doc.add_heading("最佳参数与核心指标", level=1)
    doc.add_paragraph("Best Parameters And Key Metrics")
    _add_experiment_best_section(doc, detail)

    doc.add_heading("参数稳定性", level=1)
    doc.add_paragraph("Parameter Stability")
    _add_payload_table(doc, detail.get("stability", {}), PARAMETER_STABILITY_LABELS)

    doc.add_heading("样本内/样本外验证", level=1)
    doc.add_paragraph("Train-Test Validation")
    _add_payload_table(doc, detail.get("train_test_validation", {}), VALIDATION_LABELS)

    doc.add_heading("滚动样本外验证", level=1)
    doc.add_paragraph("Walk-Forward Validation")
    _add_payload_table(doc, {key: value for key, value in dict(detail.get("walk_forward_validation", {}) or {}).items() if key != "windows"}, WALK_FORWARD_LABELS)
    _add_walk_forward_windows_table(doc, detail.get("walk_forward_validation", {}))

    doc.add_heading("实验对比图", level=1)
    doc.add_paragraph("Experiment Comparison Chart")
    _add_experiment_visuals(doc, detail["summary"])

    doc.add_heading("实验明细", level=1)
    doc.add_paragraph("Experiment Runs")
    visible_columns = _experiment_visible_columns(detail)
    _add_dataframe_table(doc, detail["summary"][visible_columns], max_rows=20)

    doc.add_heading("文件追溯", level=1)
    doc.add_paragraph("File Traceability")
    _add_payload_table(
        doc,
        {
            "summary_path": detail.get("summary_path", ""),
            "runs_path": detail.get("runs_path", ""),
            "stability_path": detail.get("stability_path", ""),
            "validation_path": detail.get("validation_path", ""),
            "walk_forward_path": detail.get("walk_forward_path", ""),
        },
        TRACEABILITY_LABELS,
    )

    _save_docx_atomic(doc, output)
    return output


def export_strategy_review_docx(
    profile_path: Path | str,
    output_path: Path | str | None = None,
    report_name: str = "StrategyReviewReport",
) -> Path:
    candidate = parse_strategy_profile_candidate(profile_path)
    strategy_dir = _candidate_strategy_dir(profile_path)
    code_source = collect_strategy_code_quality_reports(strategy_dir) if strategy_dir else collect_strategy_code_quality_reports()
    smoke_source = collect_strategy_smoke_tests(strategy_dir) if strategy_dir else collect_strategy_smoke_tests()
    code_reports = {report.strategy_id: report for report in code_source}
    smoke_reports = {report.strategy_id: report for report in smoke_source}
    approval_records = StrategyApprovalRegistry().records()
    code_report = code_reports.get(candidate.strategy_id)
    smoke_report = smoke_reports.get(candidate.strategy_id)
    readiness = evaluate_strategy_readiness_gate(candidate, code_report, approval_records, smoke_report=smoke_report)
    output = Path(output_path) if output_path is not None else unique_report_path(report_name, suffix=".docx")
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_document(doc)
    _add_strategy_review_title(doc, candidate.strategy_id)
    _add_research_notice(doc)

    doc.add_heading("执行摘要", level=1)
    doc.add_paragraph("Executive Summary")
    _add_three_column_table(
        doc,
        [
            ("策略编号", "Strategy Id", candidate.strategy_id),
            ("版本", "Version", candidate.version),
            ("综合状态", "Readiness Status", readiness.status),
            ("档案质量", "Profile Quality", f"{candidate.quality_status} / {candidate.quality_score}"),
            ("代码质量", "Code Quality", f"{code_report.status} / {code_report.score}" if code_report else "MissingCode / 0"),
            ("烟雾测试", "Smoke Test", f"{smoke_report.status} / {smoke_report.rows} rows" if smoke_report else "MissingSmoke / 0 rows"),
            ("审批状态", "Approval Status", _latest_approval_status_for_report(candidate.strategy_id, candidate.version, approval_records) or candidate.approval_status),
        ],
    )

    doc.add_heading("审批前综合门禁", level=1)
    doc.add_paragraph("Pre-Approval Readiness Gate")
    _add_three_column_table(doc, [(f"原因 {idx}", f"Reason {idx}", reason) for idx, reason in enumerate(readiness.reasons, start=1)])
    _add_three_column_table(doc, [(f"动作 {idx}", f"Action {idx}", action) for idx, action in enumerate(readiness.actions, start=1)])

    doc.add_heading("候选策略档案", level=1)
    doc.add_paragraph("Candidate Strategy Profile")
    _add_three_column_table(doc, _candidate_profile_rows(candidate))

    doc.add_heading("代码质量检查", level=1)
    doc.add_paragraph("Code Quality Check")
    _add_three_column_table(doc, _code_quality_rows(code_report))

    doc.add_heading("烟雾测试", level=1)
    doc.add_paragraph("Smoke Test")
    _add_three_column_table(doc, _smoke_test_rows(smoke_report))

    doc.add_heading("审批记录", level=1)
    doc.add_paragraph("Approval Records")
    approval_rows = _approval_record_rows(candidate.strategy_id, candidate.version, approval_records)
    if approval_rows:
        _add_three_column_table(doc, approval_rows)
    else:
        doc.add_paragraph("未找到审批记录。")
        doc.add_paragraph("No approval records were found.")

    doc.add_heading("文件追溯", level=1)
    doc.add_paragraph("File Traceability")
    _add_three_column_table(
        doc,
        [
            ("档案路径", "Profile Path", candidate.path),
            ("代码路径", "Code Path", code_report.path if code_report else ""),
            ("报告路径", "Report Path", str(output)),
        ],
    )

    _save_docx_atomic(doc, output)
    return output


def _candidate_strategy_dir(profile_path: Path | str) -> Path | None:
    path = Path(profile_path)
    sibling = path.parent.parent / "strategies"
    if sibling.exists():
        return sibling
    return None


def report_filename(report_name: str, suffix: str = ".docx", day: datetime | None = None) -> str:
    report_day = day or datetime.now()
    return f"{report_name}_{report_day.strftime('%d%m%Y')}{suffix}"


def unique_report_path(report_name: str, suffix: str = ".docx") -> Path:
    root = report_date_dir()
    candidate_name = report_name
    counter = 1
    while True:
        path = root / report_filename(candidate_name, suffix=suffix)
        try:
            fd = os.open(path, os.O_CREAT | os.O_EXCL | os.O_WRONLY, 0o644)
        except FileExistsError:
            counter += 1
            candidate_name = f"{report_name}{counter}"
            continue
        else:
            os.close(fd)
            return path


def _save_docx_atomic(doc: Document, output: Path) -> None:
    output.parent.mkdir(parents=True, exist_ok=True)
    fd, temp_name = tempfile.mkstemp(prefix=f".{output.name}.", suffix=".tmp.docx", dir=str(output.parent))
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        doc.save(temp_path)
        os.replace(temp_path, output)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def _setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.2)


def _add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PFIOS 回测研究报告")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(32, 36, 43)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("PFIOS Backtest Research Report").italic = True


def _add_experiment_title(doc: Document, experiment_name: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PFIOS 实验研究报告")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(32, 36, 43)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("PFIOS Experiment Research Report").italic = True
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.add_run(str(experiment_name)).bold = True


def _add_strategy_review_title(doc: Document, strategy_id: str) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("PFIOS 候选策略审查报告")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor(32, 36, 43)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("PFIOS Candidate Strategy Review Report").italic = True
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.add_run(str(strategy_id)).bold = True


def _add_research_notice(doc: Document) -> None:
    doc.add_paragraph("仅研究用途，禁止实盘交易，禁止真实下单。")
    doc.add_paragraph("Research only. Live trading and real order submission are prohibited.")
    doc.add_paragraph("历史回测结果可作为实盘交易前的研究参考，但不能保证未来收益。")
    doc.add_paragraph("Historical backtest results may support pre-trade research, but they do not guarantee future returns.")


def _add_executive_summary(doc: Document, result: BacktestResult) -> None:
    metrics = result.metrics
    buy_hold = _buy_and_hold_metrics(result)
    _add_executive_comparison_chart(doc, result, buy_hold)
    doc.add_paragraph("结果判读")
    doc.add_paragraph("Result Interpretation")
    judgement_rows = [
        (item.title_cn, item.title_en, f"{item.status}；{item.detail_cn}\n{item.detail_en}")
        for item in backtest_result_judgements(metrics, buy_hold)
    ]
    _add_three_column_table(doc, judgement_rows)
    rows = [
        ("策略总收益", "Strategy Total Return", _format_metric(metrics.get("total_return", 0), "percent")),
        ("策略年化收益", "Strategy Annualized Return", _format_metric(metrics.get("annualized_return", 0), "percent")),
        ("最大回撤", "Maximum Drawdown", _format_metric(metrics.get("max_drawdown", 0), "percent")),
        ("夏普比率", "Sharpe Ratio", _format_metric(metrics.get("sharpe", 0), "number")),
        ("买入持有总收益", "Buy And Hold Total Return", _format_metric(buy_hold["buy_hold_total_return"], "percent")),
        ("买入持有年化收益", "Buy And Hold Annualized Return", _format_metric(buy_hold["buy_hold_annualized_return"], "percent")),
        ("买入持有最大回撤", "Buy And Hold Max Drawdown", _format_metric(buy_hold["buy_hold_max_drawdown"], "percent")),
    ]
    _add_three_column_table(doc, rows)


def _add_executive_comparison_chart(doc: Document, result: BacktestResult, buy_hold: dict[str, float]) -> None:
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        doc.add_paragraph(f"执行摘要图表生成失败：{exc}")
        doc.add_paragraph(f"Executive summary chart generation failed: {exc}")
        return

    strategy_total = float(result.metrics.get("total_return", 0.0))
    strategy_annualized = float(result.metrics.get("annualized_return", 0.0))
    values = [
        strategy_total,
        buy_hold["buy_hold_total_return"],
        strategy_annualized,
        buy_hold["buy_hold_annualized_return"],
    ]
    labels = ["Strategy Total", "Buy Hold Total", "Strategy Annual", "Buy Hold Annual"]
    colors = ["#8A1538", "#64748B", "#0F766E", "#94A3B8"]

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "ExecutiveReturnComparison.png"
        fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=160)
        bars = ax.bar(labels, [value * 100 for value in values], color=colors)
        ax.axhline(0, color="#111827", linewidth=0.8)
        ax.set_title("Strategy Return Vs Buy And Hold")
        ax.set_ylabel("Return (%)")
        ax.grid(axis="y", alpha=0.25)
        for bar, value in zip(bars, values):
            y = bar.get_height()
            va = "bottom" if y >= 0 else "top"
            offset = 0.6 if y >= 0 else -0.6
            ax.text(bar.get_x() + bar.get_width() / 2, y + offset, f"{value:.2%}", ha="center", va=va, fontsize=8)
        fig.tight_layout()
        fig.savefig(path)
        plt.close(fig)
        _add_picture(doc, path, "策略收益与买入持有收益对比", "Strategy Return Vs Buy And Hold")


def _add_strategy_section(doc: Document, result: BacktestResult) -> None:
    strategy = result.metadata.get("strategy", {})
    approval_status = "Approved" if StrategyApprovalRegistry().is_approved(_StrategyMetadataProxy(strategy)) else "Not Approved"
    rows = [
        ("策略编号", "Strategy Id", strategy.get("strategy_id", "")),
        ("策略版本", "Strategy Version", strategy.get("version", "")),
        ("策略描述", "Strategy Description", strategy.get("description", "")),
        ("策略参数", "Strategy Parameters", json.dumps(strategy.get("params", {}), ensure_ascii=False, default=str)),
        ("策略审批状态", "Strategy Approval Status", approval_status),
    ]
    custom_spec = _custom_strategy_spec(strategy)
    if custom_spec:
        rows.extend(
            [
                ("自定义策略名称", "Custom Strategy Name", f"{custom_spec.get('display_name', '')}\n{custom_spec.get('display_name_en', '')}"),
                ("自定义策略逻辑", "Custom Strategy Logic", str(custom_spec.get("logic_key", ""))),
                ("自定义指标组合", "Custom Indicator Combination", ", ".join(custom_spec.get("indicator_keys", []))),
                ("自定义参数设置", "Custom Parameter Settings", custom_spec.get("parameter_notes", "")),
            ]
        )
    _add_three_column_table(doc, rows)
    doc.add_heading("策略研究审查", level=2)
    doc.add_paragraph("Strategy Research Review")
    _add_three_column_table(doc, _strategy_review_rows(result, strategy, approval_status))
    doc.add_heading("策略收益来源", level=2)
    doc.add_paragraph("Strategy Return Sources")
    _add_three_column_table(doc, _return_source_rows(strategy))


def _strategy_review_rows(result: BacktestResult, strategy: dict, approval_status: str) -> list[tuple[str, str, str]]:
    metrics = result.metrics
    strategy_id = str(strategy.get("strategy_id", ""))
    custom_spec = _custom_strategy_spec(strategy)
    profile = get_strategy_profile(strategy_id)
    total_return = float(metrics.get("total_return", 0.0))
    max_drawdown = float(metrics.get("max_drawdown", 0.0))
    sharpe = float(metrics.get("sharpe", 0.0))
    trade_count = int(float(metrics.get("trade_count", 0)))
    cost_total = float(metrics.get("cost_total", 0.0))
    ending_equity = float(metrics.get("ending_equity", 0.0))
    cost_ratio = cost_total / ending_equity if ending_equity else 0.0
    effectiveness = "仍需复核" if total_return <= 0 else "扣除当前模型内佣金与滑点后为正"
    effectiveness_en = "Needs review" if total_return <= 0 else "Positive after modeled commission and slippage"
    earnings = f"{custom_spec.get('return_source', '')}\n{custom_spec.get('return_source_en', '')}" if custom_spec else f"{profile.earnings}\n{profile.earnings_en}"
    persistence = f"{custom_spec.get('thesis', '')}\n{custom_spec.get('thesis_en', '')}" if custom_spec else f"{profile.persistence}\n{profile.persistence_en}"
    failure = f"{custom_spec.get('failure', '')}\n{custom_spec.get('failure_en', '')}" if custom_spec else f"{profile.failure}\n{profile.failure_en}"
    return [
        ("我赚的是什么钱？", "What Money Does The Strategy Try To Earn?", earnings),
        ("这个规律为什么会长期存在？", "Why Might The Pattern Persist?", persistence),
        ("数据是否支持？", "Does The Data Support It?", f"样本内总收益 {total_return:.2%}，夏普 {sharpe:.2f}，交易次数 {trade_count}。"),
        ("扣除费用后是否仍有效？", "Is It Still Effective After Costs?", f"{effectiveness} / {effectiveness_en}；总交易成本 {cost_total:,.2f}，成本占期末权益 {cost_ratio:.2%}。"),
        ("最大回撤能不能接受？", "Is Maximum Drawdown Acceptable?", f"最大回撤 {max_drawdown:.2%}。是否可接受需要与你的资金规模和心理承受能力一起判断。"),
        ("什么市场环境下会失效？", "When Might It Fail?", failure),
        ("失效后系统如何停止交易？", "How Should The System Stop After Failure?", _stop_rule_text(approval_status)),
    ]


def _return_source_rows(strategy: dict) -> list[tuple[str, str, str]]:
    custom_spec = _custom_strategy_spec(strategy)
    profile = get_strategy_profile(str(strategy.get("strategy_id", "")))
    source_text = str(custom_spec.get("return_source", "")) if custom_spec else ""
    source_map = {item.source for item in RETURN_SOURCE_TAXONOMY if item.source in source_text} if custom_spec else set(profile.primary_sources)
    rows = []
    for item in RETURN_SOURCE_TAXONOMY:
        status = "Primary" if item.source in source_map else "Secondary / Needs Evidence"
        rows.append((item.source, item.source_en, f"{item.explanation} / {item.explanation_en}；{item.example} / {item.example_en}；当前策略判断：{status}"))
    return rows


def _custom_strategy_spec(strategy: dict) -> dict:
    payload = strategy.get("custom_strategy_spec", {})
    return payload if isinstance(payload, dict) else {}


def _stop_rule_text(approval_status: str) -> str:
    return (
        f"当前系统只研究不实盘，审批状态为 {approval_status}。研究层面建议触发停止研究/暂停使用条件：策略未审批、数据质量非 Pass、交叉校验非 Pass、回撤超过预设阈值、费用后收益转负或样本外表现持续恶化。"
    )


def _add_decision_quality_section(doc: Document, result: DecisionQualityResult) -> None:
    doc.add_paragraph("本节用于判断一次研究结论是否完整，输出研究状态而不是实盘操作指令。")
    doc.add_paragraph(
        "This section evaluates whether the research conclusion is complete. It outputs research status, not live trading instructions."
    )
    status_note = {
        "ContinueResearch": "继续研究：当前证据较完整，但仍需持续复核未来数据、成本和样本外表现。",
        "WatchOnly": "仅观察：可以保留为观察线索，但暂不升级为高置信参考。",
        "NeedsMoreEvidence": "需要更多证据：关键验证缺失，只能用于研究复盘。",
        "DoNotUse": "暂停使用：风险门禁未通过，先修复问题再复核。",
    }.get(result.status, result.status)
    rows = [
        ("研究状态", "Research Status", result.status),
        ("状态解释", "Status Note", status_note),
        ("决策质量分", "Decision Quality Score", f"{result.score}/100"),
        ("通过项数量", "Passed Item Count", str(len(result.passed_items))),
        ("缺失证据数量", "Missing Evidence Count", str(len(result.missing_evidence))),
        ("风险提示数量", "Warning Count", str(len(result.warnings))),
    ]
    _add_three_column_table(doc, rows)

    doc.add_heading("评分维度", level=2)
    doc.add_paragraph("Score Dimensions")
    dimension_rows = [
        (item.label, item.key, f"{item.score}/10；{item.status}；{item.evidence}")
        for item in result.dimensions
    ]
    _add_three_column_table(doc, dimension_rows)

    doc.add_heading("关键缺失证据", level=2)
    doc.add_paragraph("Missing Evidence")
    if result.missing_evidence:
        _add_three_column_table(doc, [(f"缺失 {idx}", f"Missing {idx}", item) for idx, item in enumerate(result.missing_evidence, start=1)])
    else:
        doc.add_paragraph("无关键缺失证据。")
        doc.add_paragraph("No critical missing evidence.")

    doc.add_heading("下一步研究动作", level=2)
    doc.add_paragraph("Next Research Actions")
    _add_three_column_table(doc, [(f"动作 {idx}", f"Action {idx}", item) for idx, item in enumerate(result.research_actions, start=1)])

    doc.add_heading("历史模拟暴露统计", level=2)
    doc.add_paragraph("Historical Simulated Exposure")
    _add_three_column_table(doc, _decision_quality_exposure_rows(result))


def _add_risk_gate_section(doc: Document, result: RiskGateResult) -> None:
    doc.add_paragraph("本节是研究级风控结论，用来判断该策略结果是否适合继续研究、仅观察、需要补证据或暂停使用。")
    doc.add_paragraph(
        "This section is a research-level risk conclusion used to decide whether the result should continue research, stay watch-only, require more evidence, or be paused."
    )
    doc.add_paragraph("它不是实盘交易指令，不会连接券商，也不会提交真实订单。")
    doc.add_paragraph("It is not a live trading instruction, does not connect to brokers, and does not submit real orders.")
    rows = [
        ("研究状态", "Research Status", result.status),
        ("风险分数", "Risk Score", str(result.score)),
        ("触发原因数量", "Triggered Reason Count", str(len(result.reasons))),
        ("缺失证据数量", "Missing Evidence Count", str(len(result.missing_evidence))),
        ("默认最大回撤阈值", "Default Maximum Drawdown Limit", "-25.00%"),
        ("默认交易摩擦占比阈值", "Default Trading Friction Ratio Limit", "8.00%"),
    ]
    _add_three_column_table(doc, rows)
    if result.missing_evidence:
        doc.add_heading("缺失证据", level=2)
        doc.add_paragraph("Missing Evidence")
        _add_three_column_table(doc, [(f"证据 {idx}", f"Evidence {idx}", item) for idx, item in enumerate(result.missing_evidence, start=1)])
    doc.add_heading("触发原因", level=2)
    doc.add_paragraph("Triggered Reasons")
    _add_three_column_table(doc, [(f"原因 {idx}", f"Reason {idx}", reason) for idx, reason in enumerate(result.reasons, start=1)])
    doc.add_heading("建议动作", level=2)
    doc.add_paragraph("Suggested Actions")
    _add_three_column_table(doc, [(f"动作 {idx}", f"Action {idx}", action) for idx, action in enumerate(result.actions, start=1)])


def report_evidence_summary(
    result: BacktestResult,
    *,
    data_quality_report=None,
    cross_validation_result: CrossSourceValidationResult | None = None,
    risk_gate_result: RiskGateResult | None = None,
    decision_quality: DecisionQualityResult | None = None,
) -> dict[str, Any]:
    metadata = result.metadata if isinstance(result.metadata, dict) else {}
    backtest = metadata.get("backtest", {}) if isinstance(metadata.get("backtest", {}), dict) else {}
    entity = _report_entity_metadata(metadata)
    workflow = _report_workflow_metadata(metadata)
    cost_assumptions = _report_cost_assumptions(backtest)
    missing_evidence: list[str] = []

    data_quality_status = _data_quality_status(data_quality_report)
    cross_validation_status = getattr(cross_validation_result, "status", None) if cross_validation_result is not None else None
    if not data_quality_status:
        missing_evidence.append("数据质量报告")
    elif data_quality_status not in {"Pass", "Info"}:
        missing_evidence.append(f"数据质量状态需要复核：{data_quality_status}")
    if not cross_validation_status:
        missing_evidence.append("多源交叉校验")
    elif str(cross_validation_status) not in {"Pass", "Info"}:
        missing_evidence.append(f"多源交叉校验状态需要复核：{cross_validation_status}")
    if entity["entity_status"] == "MissingEntityStatus":
        missing_evidence.append("实体注册状态")
    elif entity["entity_status"] == "MissingSymbol":
        missing_evidence.append("实体缺少可分析代码")
    if not workflow["linked_request_id"] and not workflow["workflow_input_id"]:
        workflow["lineage_status"] = "ManualOrLocalOnly"
    if not cost_assumptions["complete"]:
        missing_evidence.append("完整交易成本假设")

    decision_status = decision_quality.status if decision_quality is not None else ""
    risk_status = risk_gate_result.status if risk_gate_result is not None else ""
    evidence_status = "NeedsMoreEvidence" if missing_evidence else (decision_status or risk_status or "Review")
    return {
        "schema": "PFIOSReportEvidenceV1",
        "evidence_status": evidence_status,
        "data_quality_status": data_quality_status or "Missing",
        "cross_validation_status": str(cross_validation_status or "Missing"),
        "entity": entity,
        "workflow": workflow,
        "cost_assumptions": cost_assumptions,
        "risk_gate_status": risk_status or "Missing",
        "decision_quality_status": decision_status or "Missing",
        "missing_evidence": missing_evidence,
        "conclusion_policy": _report_conclusion_policy(evidence_status, missing_evidence),
    }


def _add_report_evidence_section(doc: Document, evidence: dict[str, Any]) -> None:
    doc.add_paragraph("本节说明报告结论所依赖的证据是否完整。证据缺失时，报告结论自动降级为研究线索。")
    doc.add_paragraph("This section explains whether the evidence behind the report conclusion is complete.")
    rows = [
        ("证据状态", "Evidence Status", evidence.get("evidence_status", "")),
        ("结论策略", "Conclusion Policy", evidence.get("conclusion_policy", "")),
        ("数据质量状态", "Data Quality Status", evidence.get("data_quality_status", "")),
        ("多源校验状态", "Cross-Source Status", evidence.get("cross_validation_status", "")),
        ("实体状态", "Entity Status", evidence.get("entity", {}).get("entity_status", "")),
        ("实体代码", "Entity Symbol", evidence.get("entity", {}).get("canonical_symbol", "")),
        ("工作流来源", "Workflow Lineage", evidence.get("workflow", {}).get("lineage_status", "")),
        ("工作流输入编号", "Workflow Input Id", evidence.get("workflow", {}).get("workflow_input_id", "")),
        ("关联请求编号", "Linked Request Id", evidence.get("workflow", {}).get("linked_request_id", "")),
        ("成本假设完整", "Cost Assumptions Complete", str(evidence.get("cost_assumptions", {}).get("complete", False))),
        ("风险闸门状态", "Risk Gate Status", evidence.get("risk_gate_status", "")),
        ("决策质量状态", "Decision Quality Status", evidence.get("decision_quality_status", "")),
    ]
    _add_three_column_table(doc, rows)
    missing = list(evidence.get("missing_evidence", []) or [])
    doc.add_heading("缺失或需复核证据", level=2)
    doc.add_paragraph("Missing Or Review Evidence")
    if missing:
        _add_three_column_table(doc, [(f"证据 {idx}", f"Evidence {idx}", item) for idx, item in enumerate(missing, start=1)])
    else:
        doc.add_paragraph("未发现关键缺失证据。")
        doc.add_paragraph("No critical missing evidence was identified.")


def _report_entity_metadata(metadata: dict[str, Any]) -> dict[str, str]:
    entity = metadata.get("entity", {})
    if not isinstance(entity, dict):
        entity = {}
    registry = metadata.get("entity_registry", {})
    if not isinstance(registry, dict):
        registry = {}
    return {
        "entity_id": str(entity.get("entity_id") or registry.get("entity_id") or ""),
        "name": str(entity.get("name") or registry.get("name") or metadata.get("symbol_name") or ""),
        "market": str(entity.get("market") or registry.get("market") or metadata.get("market") or ""),
        "canonical_symbol": str(entity.get("canonical_symbol") or registry.get("canonical_symbol") or metadata.get("symbol") or ""),
        "entity_status": str(entity.get("status") or registry.get("status") or metadata.get("entity_status") or "MissingEntityStatus"),
        "source": str(entity.get("source") or registry.get("source") or "metadata"),
    }


def _report_workflow_metadata(metadata: dict[str, Any]) -> dict[str, str]:
    workflow = metadata.get("workflow", {})
    if not isinstance(workflow, dict):
        workflow = {}
    workflow_input_id = str(workflow.get("workflow_input_id") or metadata.get("workflow_input_id") or "")
    linked_request_id = str(workflow.get("linked_request_id") or metadata.get("linked_request_id") or metadata.get("request_id") or "")
    source_system = str(workflow.get("source_system") or metadata.get("source_system") or "")
    return {
        "workflow_input_id": workflow_input_id,
        "linked_request_id": linked_request_id,
        "source_system": source_system,
        "lineage_status": "Linked" if workflow_input_id or linked_request_id else "ManualOrLocalOnly",
    }


def _report_cost_assumptions(backtest: dict[str, Any]) -> dict[str, Any]:
    required = ["commission_rate", "min_commission", "slippage_bps", "market_impact_bps", "allow_short"]
    missing = [key for key in required if key not in backtest]
    return {
        "complete": not missing,
        "missing": missing,
        "commission_rate": backtest.get("commission_rate", ""),
        "min_commission": backtest.get("min_commission", ""),
        "slippage_bps": backtest.get("slippage_bps", ""),
        "market_impact_bps": backtest.get("market_impact_bps", ""),
        "allow_short": backtest.get("allow_short", ""),
    }


def _report_conclusion_policy(evidence_status: str, missing_evidence: list[str]) -> str:
    if missing_evidence or evidence_status == "NeedsMoreEvidence":
        return "NeedsMoreEvidence：证据不完整，只能作为研究线索或复盘材料。"
    if evidence_status == "DoNotUse":
        return "DoNotUse：风险门禁未通过，暂停作为研究参考。"
    if evidence_status == "WatchOnly":
        return "WatchOnly：仅观察，继续补充证据。"
    return "ContinueResearch：证据链较完整，但仍不构成实盘交易指令。"


def _data_quality_status(summary) -> str | None:
    if summary is None:
        return None
    if is_dataclass(summary):
        payload = asdict(summary)
    elif isinstance(summary, dict):
        payload = summary
    else:
        return None
    status = payload.get("quality_status")
    return str(status) if status is not None else None


def _decision_quality_exposure_rows(result: DecisionQualityResult) -> list[tuple[str, str, str]]:
    exposure = result.simulated_exposure or {}
    return [
        ("历史模拟增加暴露金额", "Simulated Exposure Increase Amount", _format_metric(_safe_float(exposure.get("simulated_exposure_increase_amount")), "currency")),
        ("历史模拟增加暴露比例", "Simulated Exposure Increase Ratio", _format_metric(_safe_float(exposure.get("simulated_exposure_increase_ratio")), "percent")),
        ("历史模拟降低暴露金额", "Simulated Exposure Reduction Amount", _format_metric(_safe_float(exposure.get("simulated_exposure_reduction_amount")), "currency")),
        ("历史模拟降低暴露比例", "Simulated Exposure Reduction Ratio", _format_metric(_safe_float(exposure.get("simulated_exposure_reduction_ratio")), "percent")),
        ("历史模拟增加暴露次数", "Simulated Exposure Increase Count", _format_metric(_safe_float(exposure.get("simulated_exposure_increase_count")), "integer")),
        ("历史模拟降低暴露次数", "Simulated Exposure Reduction Count", _format_metric(_safe_float(exposure.get("simulated_exposure_reduction_count")), "integer")),
        ("期末持仓金额", "Ending Position Value", _format_metric(_safe_float(exposure.get("ending_position_value")), "currency")),
        ("期末持仓暴露比例", "Ending Exposure Ratio", _format_metric(_safe_float(exposure.get("ending_exposure_ratio")), "percent")),
    ]


def _risk_gate_from_payload(payload: dict | None) -> RiskGateResult:
    payload = payload or {}
    return RiskGateResult(
        status=str(payload.get("status", "")),
        score=int(payload.get("score", 0) or 0),
        reasons=list(payload.get("reasons", []) or []),
        actions=list(payload.get("actions", []) or []),
        missing_evidence=list(payload.get("missing_evidence", []) or []),
    )




def _add_metrics_table(doc: Document, metrics: dict) -> None:
    rows = []
    for key, value in metrics.items():
        chinese, english, formatter = METRIC_LABELS.get(key, (key, key.replace("_", " ").title(), "number"))
        rows.append((chinese, english, _format_metric(value, formatter)))
    _add_three_column_table(doc, rows)


def _add_strategy_diagnostics_section(doc: Document, result: BacktestResult) -> None:
    diagnostics = build_strategy_diagnostics(result)
    doc.add_paragraph("本节用于判断策略交易质量、成本韧性和可能失效环境；它不是实盘交易指令。")
    doc.add_paragraph("This section assesses trade quality, cost resilience, and possible failure regimes; it is not a live trading instruction.")
    doc.add_paragraph("失效检查 Failure Checks")
    _add_dataframe_table(doc, diagnostics.failure_checks, max_rows=20)
    doc.add_paragraph("交易质量 Trade Quality")
    _add_dataframe_table(doc, _format_diagnostic_report_frame(diagnostics.trade_quality), max_rows=20)
    doc.add_paragraph("成本压力 Cost Stress")
    _add_dataframe_table(doc, _format_diagnostic_report_frame(diagnostics.cost_sensitivity), max_rows=20)
    doc.add_paragraph("市场环境分层 Market Regime Breakdown")
    _add_dataframe_table(doc, _format_diagnostic_report_frame(diagnostics.regime_breakdown), max_rows=20)
    if not diagnostics.round_trips.empty:
        doc.add_paragraph("最近完成交易回合 Recent Completed Round Trips")
        _add_dataframe_table(doc, _format_diagnostic_report_frame(diagnostics.round_trips.tail(10)), max_rows=10)
    _add_strategy_diagnostics_charts(doc, diagnostics)


def _add_strategy_diagnostics_charts(doc: Document, diagnostics) -> None:
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        doc.add_paragraph(f"策略诊断图表生成失败：{exc}")
        doc.add_paragraph(f"Strategy diagnostics chart generation failed: {exc}")
        return
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        if not diagnostics.cost_sensitivity.empty:
            _save_cost_sensitivity_chart(plt, diagnostics.cost_sensitivity, tmp_dir / "CostSensitivity.png")
            _add_picture(doc, tmp_dir / "CostSensitivity.png", "成本压力测试", "Cost Stress Test")
        if not diagnostics.regime_breakdown.empty:
            _save_regime_breakdown_chart(plt, diagnostics.regime_breakdown, tmp_dir / "RegimeBreakdown.png")
            _add_picture(doc, tmp_dir / "RegimeBreakdown.png", "市场环境收益分层", "Market Regime Return Breakdown")


def _save_cost_sensitivity_chart(plt, frame: pd.DataFrame, path: Path) -> None:
    from matplotlib.ticker import FuncFormatter

    fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=160)
    ax.plot(frame["成本倍数 Cost Multiplier"], frame["调整后总收益 Adjusted Total Return"], color="#8A1538", marker="o", linewidth=1.5)
    ax.axhline(0, color="#111827", linewidth=0.7)
    ax.set_title("Cost Stress Test")
    ax.set_ylabel("Adjusted Total Return")
    ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.0%}"))
    ax.grid(True, alpha=0.25)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _save_regime_breakdown_chart(plt, frame: pd.DataFrame, path: Path) -> None:
    from matplotlib.ticker import FuncFormatter

    fig, ax = plt.subplots(figsize=(7.2, 3.4), dpi=160)
    labels = [_english_chart_label(value) for value in frame["市场环境 Market Regime"]]
    x = range(len(labels))
    ax.bar([value - 0.18 for value in x], frame["策略收益 Strategy Return"], width=0.36, color="#8A1538", label="Strategy")
    ax.bar([value + 0.18 for value in x], frame["目标收益 Target Return"], width=0.36, color="#667085", label="Target")
    ax.axhline(0, color="#111827", linewidth=0.7)
    ax.set_title("Market Regime Return Breakdown")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels, rotation=25, ha="right", fontsize=8)
    ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.0%}"))
    ax.legend(loc="best", fontsize=8)
    ax.grid(True, axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _english_chart_label(value: object) -> str:
    text = str(value)
    if "High Volatility" in text:
        return "High Vol"
    if "Low Volatility" in text:
        return "Low Vol"
    if "Up" in text:
        return "Up"
    if "Down" in text:
        return "Down"
    if "Flat" in text:
        return "Flat"
    return text.encode("ascii", errors="ignore").decode("ascii") or "Regime"


def _format_diagnostic_report_frame(frame: pd.DataFrame) -> pd.DataFrame:
    if frame is None or frame.empty:
        return pd.DataFrame()
    display = frame.copy()
    for column in display.columns:
        if column == "Type":
            continue
        if "Return" in column or "收益" in column or "Rate" in column or "率" in column or "Worst Period" in column or "最差" in column:
            display[column] = display[column].map(lambda value: _format_metric(value, "percent"))
        elif "Cost" in column or "成本" in column or "Equity" in column or "权益" in column or "PnL" in column or "损益" in column or "盈利" in column or "亏损" in column:
            display[column] = display[column].map(lambda value: _format_metric(value, "currency"))
        elif column == "Value":
            display[column] = [
                _format_metric(value, value_type)
                for value, value_type in zip(display["Value"], display.get("Type", ["number"] * len(display)))
            ]
    if "Type" in display.columns:
        display = display.drop(columns=["Type"])
    return display


def _add_three_column_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for index, header in enumerate(["中文", "English", "Value"]):
        cell = table.rows[0].cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for chinese, english, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(chinese)
        cells[1].text = str(english)
        cells[2].text = str(value)


def _add_dataframe_table(doc: Document, frame, max_rows: int = 20) -> None:
    if frame is None or frame.empty:
        doc.add_paragraph("无记录。")
        doc.add_paragraph("No records.")
        return
    display = frame.tail(max_rows).copy()
    table = doc.add_table(rows=1, cols=len(display.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(display.columns):
        table.rows[0].cells[idx].text = str(column)
    for _, row in display.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = _format_table_value(value)


def _add_charts(doc: Document, result: BacktestResult) -> None:
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        doc.add_paragraph(f"图表生成失败：{exc}")
        doc.add_paragraph(f"Chart generation failed: {exc}")
        return

    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        equity = result.equity_curve.copy()
        equity["datetime"] = pd.to_datetime(equity["datetime"])
        _save_line_chart(plt, equity["datetime"], equity["equity"], "Equity Curve", "Equity", tmp_dir / "EquityCurve.png", "#8A1538")
        _add_picture(doc, tmp_dir / "EquityCurve.png", "收益曲线", "Equity Curve")
        _save_line_chart(plt, equity["datetime"], equity["drawdown"], "Drawdown Curve", "Drawdown", tmp_dir / "DrawdownCurve.png", "#0F766E")
        _add_picture(doc, tmp_dir / "DrawdownCurve.png", "回撤曲线", "Drawdown Curve")
        price_frame = _price_frame(result)
        if price_frame is not None and not price_frame.empty:
            _save_return_comparison_curve(plt, equity, price_frame, tmp_dir / "ReturnComparisonCurve.png")
            _add_picture(doc, tmp_dir / "ReturnComparisonCurve.png", "策略、目标与相对收益曲线", "Strategy, Target, And Relative Return Curve")
            _save_monthly_return_heatmap(plt, equity, tmp_dir / "MonthlyReturnHeatmap.png")
            _add_picture(doc, tmp_dir / "MonthlyReturnHeatmap.png", "月度收益热力图", "Monthly Return Heatmap")
            _save_rolling_risk_chart(plt, equity, tmp_dir / "RollingRisk.png")
            _add_picture(doc, tmp_dir / "RollingRisk.png", "滚动夏普与波动率", "Rolling Sharpe And Volatility")
            _save_trade_chart(plt, price_frame, result.trades, tmp_dir / "PriceAndTrades.png")
            _add_picture(doc, tmp_dir / "PriceAndTrades.png", "买卖点图", "Price And Trades")


def _save_line_chart(plt, x, y, title: str, ylabel: str, path: Path, color: str) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=160)
    ax.plot(x, y, color=color, linewidth=1.6)
    ax.set_title(title)
    ax.set_ylabel(ylabel)
    ax.grid(True, alpha=0.25)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _save_return_comparison_curve(plt, equity: pd.DataFrame, price_frame: pd.DataFrame, path: Path) -> None:
    from matplotlib.ticker import FuncFormatter

    frame = _return_comparison_frame(equity, price_frame)
    fig, ax = plt.subplots(figsize=(7.2, 3.4), dpi=160)
    ax.plot(frame["datetime"], frame["strategy_return"], color="#8A1538", linewidth=1.5, label="Strategy")
    ax.plot(frame["datetime"], frame["target_return"], color="#475467", linewidth=1.4, label="Target")
    ax.plot(frame["datetime"], frame["relative_return"], color="#2563EB", linewidth=1.2, linestyle="--", label="Relative")
    ax.axhline(0, color="#111827", linewidth=0.7)
    ax.set_title("Strategy, Target, And Relative Return")
    ax.set_ylabel("Return")
    ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.0%}"))
    ax.grid(True, alpha=0.25)
    ax.legend(loc="best", fontsize=8)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _return_comparison_frame(equity: pd.DataFrame, price_frame: pd.DataFrame) -> pd.DataFrame:
    strategy = equity[["datetime", "equity"]].copy()
    strategy["datetime"] = pd.to_datetime(strategy["datetime"])
    strategy["equity"] = pd.to_numeric(strategy["equity"], errors="coerce")
    first_equity = float(strategy["equity"].dropna().iloc[0]) if not strategy["equity"].dropna().empty else 0.0
    strategy["strategy_return"] = strategy["equity"] / first_equity - 1.0 if first_equity else 0.0

    target = price_frame[["datetime", "close"]].copy()
    target["datetime"] = pd.to_datetime(target["datetime"])
    target["close"] = pd.to_numeric(target["close"], errors="coerce")
    first_close = float(target["close"].dropna().iloc[0]) if not target["close"].dropna().empty else 0.0
    target["target_return"] = target["close"] / first_close - 1.0 if first_close else 0.0

    frame = strategy[["datetime", "strategy_return"]].merge(target[["datetime", "target_return"]], on="datetime", how="inner")
    frame["relative_return"] = frame["strategy_return"] - frame["target_return"]
    return frame.sort_values("datetime").reset_index(drop=True)


def _save_monthly_return_heatmap(plt, equity: pd.DataFrame, path: Path) -> None:
    from matplotlib.ticker import FuncFormatter

    monthly = _monthly_return_matrix(equity)
    fig, ax = plt.subplots(figsize=(7.2, 3.4), dpi=160)
    if monthly.empty:
        ax.text(0.5, 0.5, "Not enough data", ha="center", va="center")
        ax.axis("off")
    else:
        values = monthly.to_numpy(dtype=float)
        clean_values = pd.Series(values.flatten()).dropna()
        limit = max(float(clean_values.abs().max()), 0.01) if not clean_values.empty else 0.01
        image = ax.imshow(values, cmap="RdYlGn_r", vmin=-limit, vmax=limit, aspect="auto")
        ax.set_title("Monthly Return Heatmap")
        ax.set_xticks(range(len(monthly.columns)))
        ax.set_xticklabels([str(column) for column in monthly.columns], fontsize=8)
        ax.set_yticks(range(len(monthly.index)))
        ax.set_yticklabels([str(index) for index in monthly.index], fontsize=8)
        for y in range(values.shape[0]):
            for x in range(values.shape[1]):
                value = values[y, x]
                if pd.notna(value):
                    ax.text(x, y, f"{value:.1%}", ha="center", va="center", fontsize=6, color="#111827")
        fig.colorbar(image, ax=ax, fraction=0.025, pad=0.02, format=FuncFormatter(lambda value, _: f"{value:.0%}"))
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _monthly_return_matrix(equity: pd.DataFrame) -> pd.DataFrame:
    frame = equity[["datetime", "equity"]].copy()
    frame["datetime"] = pd.to_datetime(frame["datetime"])
    frame["equity"] = pd.to_numeric(frame["equity"], errors="coerce")
    frame = frame.dropna().set_index("datetime").sort_index()
    if len(frame) < 2:
        return pd.DataFrame()
    monthly_equity = frame["equity"].resample("ME").last().dropna()
    if len(monthly_equity) < 2:
        monthly_equity = frame["equity"].resample("M").last().dropna()
    monthly_return = monthly_equity.pct_change().dropna()
    if monthly_return.empty:
        return pd.DataFrame()
    table = monthly_return.to_frame("return")
    table["year"] = table.index.year
    table["month"] = table.index.month
    return table.pivot(index="year", columns="month", values="return").reindex(columns=range(1, 13))


def _save_rolling_risk_chart(plt, equity: pd.DataFrame, path: Path, window: int = 63) -> None:
    from matplotlib.ticker import FuncFormatter

    frame = _rolling_risk_frame(equity, window=window)
    fig, ax1 = plt.subplots(figsize=(7.2, 3.4), dpi=160)
    if frame.empty:
        ax1.text(0.5, 0.5, "Not enough data", ha="center", va="center")
        ax1.axis("off")
    else:
        ax1.plot(frame["datetime"], frame["rolling_sharpe"], color="#8A1538", linewidth=1.3, label="Rolling Sharpe")
        ax1.axhline(0, color="#111827", linewidth=0.7)
        ax1.set_ylabel("Sharpe")
        ax2 = ax1.twinx()
        ax2.plot(frame["datetime"], frame["rolling_volatility"], color="#0F766E", linewidth=1.2, label="Rolling Volatility")
        ax2.set_ylabel("Volatility")
        ax2.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.0%}"))
        ax1.set_title(f"Rolling Sharpe And Volatility ({window} bars)")
        ax1.grid(True, alpha=0.25)
        lines = ax1.get_lines() + ax2.get_lines()
        ax1.legend(lines, [line.get_label() for line in lines], loc="best", fontsize=8)
        fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _rolling_risk_frame(equity: pd.DataFrame, window: int = 63) -> pd.DataFrame:
    frame = equity[["datetime", "equity"]].copy()
    frame["datetime"] = pd.to_datetime(frame["datetime"])
    frame["equity"] = pd.to_numeric(frame["equity"], errors="coerce")
    frame = frame.dropna().sort_values("datetime")
    if len(frame) < max(5, window // 2):
        return pd.DataFrame()
    returns = frame["equity"].pct_change().fillna(0.0)
    rolling_mean = returns.rolling(window, min_periods=max(5, window // 3)).mean()
    rolling_std = returns.rolling(window, min_periods=max(5, window // 3)).std()
    rolling_volatility = rolling_std * 252**0.5
    rolling_sharpe = (rolling_mean * 252 / rolling_volatility).replace([float("inf"), -float("inf")], 0.0).fillna(0.0)
    return pd.DataFrame(
        {
            "datetime": frame["datetime"],
            "rolling_sharpe": rolling_sharpe,
            "rolling_volatility": rolling_volatility.fillna(0.0),
        }
    )


def _save_trade_chart(plt, price_frame: pd.DataFrame, trades: pd.DataFrame, path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=160)
    ax.plot(price_frame["datetime"], price_frame["close"], color="#1F2937", linewidth=1.2, label="Close")
    if trades is not None and not trades.empty:
        trades = trades.copy()
        trades["datetime"] = pd.to_datetime(trades["datetime"])
        buys = trades[trades["side"] == "BUY"]
        sells = trades[trades["side"] == "SELL"]
        ax.scatter(
            buys["datetime"],
            buys["price"],
            facecolors="none",
            edgecolors="#D97706",
            linewidths=0.9,
            s=18,
            label="Buy",
            zorder=3,
        )
        ax.scatter(
            sells["datetime"],
            sells["price"],
            color="#2563EB",
            marker="x",
            linewidths=0.9,
            s=18,
            label="Sell",
            zorder=3,
        )
    ax.set_title("Price And Trades")
    ax.set_ylabel("Price")
    ax.legend(loc="best")
    ax.grid(True, alpha=0.25)
    fig.autofmt_xdate()
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _add_bootstrap_robustness_section(doc: Document, result: BacktestResult, target_return: float = 0.0) -> None:
    robustness = bootstrap_equity_robustness(
        result.equity_curve,
        simulations=_report_bootstrap_simulations(),
        seed=42,
        target_return=float(target_return),
    )
    _add_dataframe_table(doc, pd.DataFrame(robustness_summary_rows(robustness)), max_rows=20)
    doc.add_paragraph("Bootstrap 使用历史策略收益重采样生成模拟路径，用于检查结果脆弱性；它不是未来收益预测。")
    doc.add_paragraph("Bootstrap resamples historical strategy returns to generate simulated paths for fragility checks; it is not a forecast.")
    if robustness.simulations.empty:
        return
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        doc.add_paragraph(f"Bootstrap 图表生成失败：{exc}")
        doc.add_paragraph(f"Bootstrap chart generation failed: {exc}")
        return
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        _save_bootstrap_distribution_chart(
            plt,
            robustness.simulations,
            "total_return",
            "Simulated Total Return Distribution",
            "Total Return",
            tmp_dir / "BootstrapTotalReturn.png",
            "#8A1538",
            "percent",
        )
        _add_picture(doc, tmp_dir / "BootstrapTotalReturn.png", "模拟总收益分布", "Simulated Total Return Distribution")
        _save_bootstrap_distribution_chart(
            plt,
            robustness.simulations,
            "max_drawdown",
            "Simulated Max Drawdown Distribution",
            "Max Drawdown",
            tmp_dir / "BootstrapMaxDrawdown.png",
            "#0F766E",
            "percent",
        )
        _add_picture(doc, tmp_dir / "BootstrapMaxDrawdown.png", "模拟最大回撤分布", "Simulated Max Drawdown Distribution")
        if not robustness.sample_paths.empty:
            _save_bootstrap_paths_chart(plt, robustness.sample_paths, robustness.path_interval, tmp_dir / "BootstrapPaths.png")
            _add_picture(doc, tmp_dir / "BootstrapPaths.png", "Bootstrap 样本收益路径", "Bootstrap Sample Return Paths")


def _report_bootstrap_simulations() -> int:
    value = os.environ.get("PFI_TEST_BOOTSTRAP_SIMULATIONS")
    if value:
        try:
            return max(10_000, int(value))
        except ValueError:
            return 10_000
    return 10_000


def _save_bootstrap_distribution_chart(
    plt,
    simulations: pd.DataFrame,
    column: str,
    title: str,
    xlabel: str,
    path: Path,
    color: str,
    formatter: str,
) -> None:
    from matplotlib.ticker import FuncFormatter

    fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=160)
    values = pd.to_numeric(simulations[column], errors="coerce").dropna()
    ax.hist(values, bins=32, color=color, alpha=0.84)
    ax.axvline(values.median(), color="#111827", linewidth=1.0, linestyle="--", label="Median")
    ax.set_title(title)
    ax.set_xlabel(xlabel)
    ax.set_ylabel("Count")
    if formatter == "percent":
        ax.xaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.0%}"))
    ax.grid(axis="y", alpha=0.22)
    ax.legend(loc="best", fontsize=8)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _save_bootstrap_paths_chart(plt, sample_paths: pd.DataFrame, path_interval: pd.DataFrame, path: Path) -> None:
    from matplotlib.ticker import FuncFormatter

    fig, ax = plt.subplots(figsize=(7.2, 3.2), dpi=160)
    if path_interval is not None and not path_interval.empty:
        ax.fill_between(
            path_interval["step"],
            path_interval["p05"] - 1.0,
            path_interval["p95"] - 1.0,
            color="#7DD3FC",
            alpha=0.24,
            label="5%-95% Band",
        )
        ax.plot(path_interval["step"], path_interval["median"] - 1.0, color="#0EA5E9", linewidth=1.4, label="Median")
    for column in [column for column in sample_paths.columns if column != "step"]:
        ax.plot(sample_paths["step"], sample_paths[column] - 1.0, color="#8A1538", alpha=0.16, linewidth=0.8)
    ax.axhline(0, color="#111827", linewidth=0.7)
    ax.set_title("Bootstrap Sample Return Paths")
    ax.set_xlabel("Step")
    ax.set_ylabel("Return")
    ax.yaxis.set_major_formatter(FuncFormatter(lambda value, _: f"{value:.0%}"))
    ax.grid(True, alpha=0.22)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _add_picture(doc: Document, path: Path, title_cn: str, title_en: str) -> None:
    caption = doc.add_paragraph()
    caption.add_run(title_cn).bold = True
    doc.add_paragraph(title_en)
    doc.add_picture(str(path), width=Inches(6.6))


def _price_frame(result: BacktestResult) -> pd.DataFrame | None:
    if result.positions.empty:
        return None
    positions = result.positions.copy()
    if "symbol" in positions.columns and positions["symbol"].nunique() > 1:
        return None
    positions["datetime"] = pd.to_datetime(positions["datetime"])
    return positions[["datetime", "close"]].drop_duplicates("datetime").sort_values("datetime")


def _buy_and_hold_metrics(result: BacktestResult) -> dict[str, float]:
    prices = _price_frame(result)
    if prices is None or prices.empty or len(prices) < 2:
        return buy_and_hold_metrics(pd.DataFrame(), annualization=252)
    return buy_and_hold_metrics(prices, annualization=252)


def _add_buy_and_hold_explanation(doc: Document, buy_hold_metrics: dict[str, float]) -> None:
    doc.add_paragraph(
        "买入持有基准表示在回测第一根可用价格买入标的并一直持有到最后一根价格，不做任何策略调仓。"
    )
    doc.add_paragraph(
        "The buy-and-hold benchmark means buying the symbol at the first available price and holding it to the last available price without strategy rebalancing."
    )
    doc.add_paragraph("公式：买入持有总收益率 = 期末收盘价 / 期初收盘价 - 1。")
    doc.add_paragraph("Formula: Buy And Hold Total Return = Ending Close / Starting Close - 1.")
    doc.add_paragraph("公式：买入持有年化收益率 = (1 + 总收益率) ^ (252 / 价格序列期数) - 1。")
    doc.add_paragraph("Formula: Buy And Hold Annualized Return = (1 + Total Return) ^ (252 / Price Periods) - 1.")
    doc.add_paragraph("公式：买入持有最大回撤 = 买入持有财富曲线 / 历史最高财富 - 1 的最小值。")
    doc.add_paragraph("Formula: Buy And Hold Max Drawdown = the minimum of buy-and-hold wealth / running peak wealth - 1.")


def _add_data_quality_summary(doc: Document, summary) -> None:
    if summary is None:
        doc.add_paragraph("未提供数据质量报告。")
        doc.add_paragraph("No data quality report was provided.")
        return
    if is_dataclass(summary):
        payload = asdict(summary)
    elif isinstance(summary, dict):
        payload = summary
    else:
        payload = {"summary": str(summary)}
    rows = []
    for key, value in payload.items():
        chinese, english = DATA_QUALITY_LABELS.get(key, (key, key.replace("_", " ").title()))
        rows.append((chinese, english, _format_table_value(value)))
    _add_three_column_table(doc, rows)


def _add_cross_validation_summary(doc: Document, result: CrossSourceValidationResult | None) -> None:
    if result is None:
        doc.add_paragraph("未提供多源交叉校验结果。")
        doc.add_paragraph("No cross-source validation result was provided.")
        return
    rows = [
        ("校验状态", "Validation Status", result.status),
        ("标的", "Symbol", result.symbol),
        ("市场", "Market", result.market),
        ("周期", "Interval", result.interval),
        ("数据源", "Providers", ", ".join(result.providers)),
        ("重叠行数", "Overlap Rows", str(result.overlap_rows)),
        ("最大收盘价差异", "Max Close Difference", _format_metric(result.max_close_diff_pct, "percent")),
        ("平均收盘价差异", "Mean Close Difference", _format_metric(result.mean_close_diff_pct, "percent")),
    ]
    _add_three_column_table(doc, rows)


def _candidate_profile_rows(candidate) -> list[tuple[str, str, str]]:
    return [
        ("策略编号", "Strategy Id", candidate.strategy_id),
        ("中文名称", "Chinese Name", candidate.display_name),
        ("英文名称", "English Name", candidate.display_name_en),
        ("版本", "Version", candidate.version),
        ("类别", "Category", candidate.category),
        ("研究假设", "Research Thesis", candidate.thesis),
        ("收益来源", "Return Source", candidate.return_source),
        ("失效环境", "Failure Regime", candidate.failure),
        ("参数设置", "Parameter Settings", candidate.parameter_notes),
        ("质量状态", "Quality Status", candidate.quality_status),
        ("质量分数", "Quality Score", str(candidate.quality_score)),
        ("缺失项", "Missing Items", ", ".join(candidate.missing_items)),
    ]


def _code_quality_rows(report) -> list[tuple[str, str, str]]:
    if report is None:
        return [("代码状态", "Code Status", "MissingCode"), ("代码分数", "Code Score", "0"), ("发现项", "Findings", "No matching custom strategy code file.")]
    return [
        ("策略编号", "Strategy Id", report.strategy_id),
        ("代码状态", "Code Status", report.status),
        ("代码分数", "Code Score", str(report.score)),
        ("发现项", "Findings", ", ".join(report.findings) if report.findings else "No findings."),
        ("代码路径", "Code Path", report.path),
    ]


def _smoke_test_rows(report) -> list[tuple[str, str, str]]:
    if report is None:
        return [("烟雾测试", "Smoke Test", "MissingSmoke"), ("信号行数", "Signal Rows", "0"), ("发现项", "Findings", "No smoke test report was available.")]
    return [
        ("策略编号", "Strategy Id", report.strategy_id),
        ("烟雾测试", "Smoke Test", report.status),
        ("信号行数", "Signal Rows", str(report.rows)),
        ("发现项", "Findings", ", ".join(report.findings) if report.findings else "No findings."),
        ("代码路径", "Code Path", report.path),
    ]


def _approval_record_rows(strategy_id: str, version: str, records) -> list[tuple[str, str, str]]:
    rows = []
    for index, record in enumerate(records, start=1):
        if record.strategy_id == strategy_id and record.version == version:
            rows.extend(
                [
                    (f"审批记录 {index}", f"Approval Record {index}", record.approval_id),
                    ("审批状态", "Approval Status", record.status),
                    ("修改说明", "Change Summary", record.change_summary),
                    ("风险说明", "Risk Notes", record.risk_notes),
                    ("申请时间", "Requested At", record.requested_at),
                    ("批准时间", "Approved At", record.approved_at or ""),
                ]
            )
    return rows


def _latest_approval_status_for_report(strategy_id: str, version: str, records) -> str:
    matches = [record for record in records if record.strategy_id == strategy_id and record.version == version]
    return matches[-1].status if matches else ""


def _add_experiment_executive_summary(doc: Document, detail: dict[str, object]) -> None:
    best = dict(detail.get("best_run", {}) or {})
    stability = dict(detail.get("stability", {}) or {})
    train_test = dict(detail.get("train_test_validation", {}) or {})
    walk_forward = dict(detail.get("walk_forward_validation", {}) or {})
    risk_gate = dict(detail.get("risk_gate", {}) or {})
    _add_experiment_chart(doc, detail["summary"])
    rows = [
        ("实验名称", "Experiment Name", str(detail.get("experiment", ""))),
        ("运行数量", "Run Count", str(detail.get("run_count", ""))),
        ("最佳运行编号", "Best Run Id", str(best.get("run_id", ""))),
        ("最佳总收益", "Best Total Return", _format_metric(_safe_float(best.get("total_return")), "percent")),
        ("最佳夏普比率", "Best Sharpe Ratio", _format_metric(_safe_float(best.get("sharpe")), "number")),
        ("最佳最大回撤", "Best Maximum Drawdown", _format_metric(_safe_float(best.get("max_drawdown")), "percent")),
        ("参数稳定性状态", "Parameter Stability Status", str(stability.get("stability_status", ""))),
        ("样本外验证状态", "Train-Test Validation Status", str(train_test.get("validation_status", ""))),
        ("滚动验证状态", "Walk-Forward Validation Status", str(walk_forward.get("validation_status", ""))),
        ("研究状态", "Research Status", str(risk_gate.get("status", ""))),
    ]
    _add_three_column_table(doc, rows)


def _add_experiment_best_section(doc: Document, detail: dict[str, object]) -> None:
    best = dict(detail.get("best_run", {}) or {})
    params = dict(detail.get("best_params", {}) or {})
    doc.add_heading("最佳参数", level=2)
    doc.add_paragraph("Best Parameters")
    if params:
        _add_three_column_table(doc, [(str(key), str(key).replace("_", " ").title(), _format_table_value(value)) for key, value in params.items()])
    else:
        doc.add_paragraph("未识别最佳参数。")
        doc.add_paragraph("No best parameters were identified.")
    doc.add_heading("最佳运行指标", level=2)
    doc.add_paragraph("Best Run Metrics")
    metric_rows = []
    for key in detail.get("metric_columns", []):
        chinese, english, formatter = METRIC_LABELS.get(str(key), (str(key), str(key).replace("_", " ").title(), "number"))
        metric_rows.append((chinese, english, _format_metric(_safe_float(best.get(key)), formatter)))
    _add_three_column_table(doc, metric_rows)


def _add_payload_table(doc: Document, payload: dict | None, labels: dict[str, tuple[str, str]]) -> None:
    payload = payload or {}
    if not payload:
        doc.add_paragraph("未提供记录。")
        doc.add_paragraph("No record was provided.")
        return
    rows = []
    for key, value in payload.items():
        if isinstance(value, (list, dict)):
            value = json.dumps(value, ensure_ascii=False, default=str)
        chinese, english = labels.get(str(key), (str(key), str(key).replace("_", " ").title()))
        rows.append((chinese, english, _format_table_value(value)))
    _add_three_column_table(doc, rows)


def _add_walk_forward_windows_table(doc: Document, validation: dict | None) -> None:
    windows = dict(validation or {}).get("windows", [])
    if not windows:
        return
    doc.add_heading("滚动窗口明细", level=2)
    doc.add_paragraph("Walk-Forward Window Details")
    frame = pd.DataFrame(windows)
    if "best_params" in frame.columns:
        frame["best_params"] = frame["best_params"].map(lambda value: json.dumps(value, ensure_ascii=False, default=str))
    _add_dataframe_table(doc, frame, max_rows=20)


def _add_experiment_chart(doc: Document, summary: pd.DataFrame) -> None:
    if summary is None or summary.empty:
        doc.add_paragraph("实验对比图无法生成：实验明细为空。")
        doc.add_paragraph("Experiment comparison chart cannot be generated because the experiment summary is empty.")
        return
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        doc.add_paragraph(f"实验对比图生成失败：{exc}")
        doc.add_paragraph(f"Experiment comparison chart generation failed: {exc}")
        return
    display = summary.head(10).copy()
    labels = display["run_id"].astype(str).tolist() if "run_id" in display.columns else [str(idx + 1) for idx in range(len(display))]
    total_source = display["total_return"] if "total_return" in display.columns else pd.Series([0.0] * len(display))
    sharpe_source = display["sharpe"] if "sharpe" in display.columns else pd.Series([0.0] * len(display))
    total_return = pd.to_numeric(total_source, errors="coerce").fillna(0) * 100
    sharpe = pd.to_numeric(sharpe_source, errors="coerce").fillna(0)
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "ExperimentComparison.png"
        fig, ax1 = plt.subplots(figsize=(7.2, 3.4), dpi=160)
        x = range(len(display))
        ax1.bar(x, total_return, color="#8A1538", alpha=0.82, label="Total Return (%)")
        ax1.set_ylabel("Total Return (%)")
        ax1.set_xticks(list(x))
        ax1.set_xticklabels(labels, rotation=35, ha="right", fontsize=7)
        ax2 = ax1.twinx()
        ax2.plot(list(x), sharpe, color="#0F766E", marker="o", linewidth=1.4, label="Sharpe")
        ax2.set_ylabel("Sharpe")
        ax1.set_title("Top Experiment Runs")
        ax1.grid(axis="y", alpha=0.22)
        fig.tight_layout()
        fig.savefig(path)
        plt.close(fig)
        _add_picture(doc, path, "实验 Top Runs 对比", "Top Experiment Runs Comparison")


def _add_experiment_visuals(doc: Document, summary: pd.DataFrame) -> None:
    _add_experiment_chart(doc, summary)
    if summary is None or summary.empty:
        return
    try:
        import matplotlib

        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except Exception as exc:
        doc.add_paragraph(f"参数热力图生成失败：{exc}")
        doc.add_paragraph(f"Parameter heatmap generation failed: {exc}")
        return
    param_columns = [column for column in summary.columns if column.startswith("param_")]
    if len(param_columns) < 2:
        doc.add_paragraph("参数热力图需要至少两个参数维度。")
        doc.add_paragraph("Parameter heatmaps require at least two parameter dimensions.")
        return
    with tempfile.TemporaryDirectory() as tmp:
        tmp_dir = Path(tmp)
        specs = [
            ("total_return", "总收益参数热力图", "Total Return Parameter Heatmap", "RdYlGn_r", "percent"),
            ("sharpe", "夏普参数热力图", "Sharpe Parameter Heatmap", "viridis", "number"),
            ("max_drawdown", "最大回撤参数热力图", "Max Drawdown Parameter Heatmap", "RdYlGn_r", "percent"),
        ]
        for metric, title_cn, title_en, cmap, formatter in specs:
            heatmap = _experiment_parameter_heatmap_frame(summary, metric)
            if heatmap.empty:
                continue
            path = tmp_dir / f"{metric}_ParameterHeatmap.png"
            _save_experiment_parameter_heatmap(plt, heatmap, metric, title_en, cmap, formatter, path)
            _add_picture(doc, path, title_cn, title_en)


def _experiment_parameter_heatmap_frame(summary: pd.DataFrame, metric: str) -> pd.DataFrame:
    param_columns = [column for column in summary.columns if column.startswith("param_")]
    if len(param_columns) < 2 or metric not in summary.columns:
        return pd.DataFrame()
    x_param, y_param = param_columns[:2]
    frame = summary[[x_param, y_param, metric]].copy()
    frame[metric] = pd.to_numeric(frame[metric], errors="coerce")
    frame = frame.dropna(subset=[metric])
    if frame.empty:
        return pd.DataFrame()
    return frame.pivot_table(index=y_param, columns=x_param, values=metric, aggfunc="mean").sort_index().sort_index(axis=1)


def _save_experiment_parameter_heatmap(plt, heatmap: pd.DataFrame, metric: str, title: str, cmap: str, formatter: str, path: Path) -> None:
    from matplotlib.ticker import FuncFormatter

    values = heatmap.to_numpy(dtype=float)
    clean_values = pd.Series(values.flatten()).dropna()
    fig, ax = plt.subplots(figsize=(7.2, 3.4), dpi=160)
    if clean_values.empty:
        ax.text(0.5, 0.5, "No data", ha="center", va="center")
        ax.axis("off")
    else:
        if formatter == "percent":
            limit = max(float(clean_values.abs().max()), 0.01)
            vmin, vmax = -limit, limit
        else:
            vmin, vmax = float(clean_values.min()), float(clean_values.max())
            if vmin == vmax:
                vmin -= 1.0
                vmax += 1.0
        image = ax.imshow(values, cmap=cmap, vmin=vmin, vmax=vmax, aspect="auto")
        ax.set_title(title)
        ax.set_xlabel(str(heatmap.columns.name or "Parameter X").replace("param_", ""))
        ax.set_ylabel(str(heatmap.index.name or "Parameter Y").replace("param_", ""))
        ax.set_xticks(range(len(heatmap.columns)))
        ax.set_xticklabels([str(column) for column in heatmap.columns], fontsize=8)
        ax.set_yticks(range(len(heatmap.index)))
        ax.set_yticklabels([str(index) for index in heatmap.index], fontsize=8)
        for y in range(values.shape[0]):
            for x in range(values.shape[1]):
                value = values[y, x]
                if pd.notna(value):
                    label = f"{value:.1%}" if formatter == "percent" else f"{value:.2f}"
                    ax.text(x, y, label, ha="center", va="center", fontsize=7, color="#111827")
        color_formatter = FuncFormatter(lambda value, _: f"{value:.0%}") if formatter == "percent" else None
        fig.colorbar(image, ax=ax, fraction=0.025, pad=0.02, format=color_formatter)
    fig.tight_layout()
    fig.savefig(path)
    plt.close(fig)


def _experiment_visible_columns(detail: dict[str, object]) -> list[str]:
    summary = detail["summary"]
    metric_columns = ["run_id"] + [column for column in detail.get("metric_columns", []) if column in summary.columns]
    parameter_columns = [column for column in summary.columns if column.startswith("param_")]
    return metric_columns + parameter_columns


def _add_metadata_explanation(doc: Document, result: BacktestResult) -> None:
    doc.add_paragraph(
        "运行配置与追溯信息用于说明本次结果是如何生成的，方便以后复核策略版本、参数、资金、交易成本和风控假设。"
    )
    doc.add_paragraph(
        "Run configuration and traceability explain how this result was generated, so the strategy version, parameters, capital, trading costs, and risk assumptions can be reviewed later."
    )
    strategy = result.metadata.get("strategy", {})
    backtest = result.metadata.get("backtest", {})
    rows = [
        ("策略编号", "Strategy Id", strategy.get("strategy_id", "")),
        ("策略版本", "Strategy Version", strategy.get("version", "")),
        ("策略参数", "Strategy Parameters", json.dumps(strategy.get("params", {}), ensure_ascii=False, default=str)),
        ("初始资金", "Initial Cash", _format_metric(float(backtest.get("initial_cash", 0)), "currency")),
        ("佣金率", "Commission Rate", _format_metric(float(backtest.get("commission_rate", 0)), "percent")),
        ("最低佣金", "Minimum Commission", _format_metric(float(backtest.get("min_commission", 0)), "currency")),
        ("滑点基点", "Slippage Bps", str(backtest.get("slippage_bps", ""))),
        ("冲击成本基点", "Market Impact Bps", str(backtest.get("market_impact_bps", ""))),
        ("允许做空", "Allow Short", str(backtest.get("allow_short", ""))),
    ]
    _add_three_column_table(doc, rows)


class _StrategyMetadataProxy:
    def __init__(self, metadata: dict):
        self.strategy_id = metadata.get("strategy_id", "")
        self.version = metadata.get("version", "")


def _format_metric(value: object, formatter: str = "number") -> str:
    if isinstance(value, float) and not math.isfinite(value):
        return "N/A"
    if isinstance(value, float) and formatter == "percent":
        return f"{value:.2%}"
    if isinstance(value, float) and formatter == "currency":
        return f"{value:,.2f}"
    if isinstance(value, float) and formatter == "integer":
        return f"{int(value)}"
    if isinstance(value, float):
        return f"{value:.2f}"
    return str(value)


def _format_table_value(value: object) -> str:
    if isinstance(value, float):
        return f"{value:,.2f}"
    return str(value)


def _safe_float(value: object) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return 0.0


METRIC_LABELS = {
    "total_return": ("总收益", "Total Return", "percent"),
    "annualized_return": ("年化收益", "Annualized Return", "percent"),
    "volatility": ("波动率", "Volatility", "percent"),
    "sharpe": ("夏普比率", "Sharpe Ratio", "number"),
    "sortino": ("索提诺比率", "Sortino Ratio", "number"),
    "calmar": ("卡玛比率", "Calmar Ratio", "number"),
    "max_drawdown": ("最大回撤", "Maximum Drawdown", "percent"),
    "win_rate": ("胜率", "Win Rate", "percent"),
    "trade_count": ("交易次数", "Trade Count", "integer"),
    "buy_count": ("买入次数", "Buy Count", "integer"),
    "sell_count": ("卖出次数", "Sell Count", "integer"),
    "round_trip_count": ("完整交易次数", "Round Trip Count", "integer"),
    "turnover": ("换手率", "Turnover", "number"),
    "average_gain": ("平均盈利", "Average Gain", "currency"),
    "average_loss": ("平均亏损", "Average Loss", "currency"),
    "cost_total": ("总建模交易摩擦", "Total Modeled Trading Friction", "currency"),
    "ending_equity": ("期末权益", "Ending Equity", "currency"),
    "buy_hold_total_return": ("买入持有总收益", "Buy And Hold Total Return", "percent"),
    "buy_hold_annualized_return": ("买入持有年化收益", "Buy And Hold Annualized Return", "percent"),
    "buy_hold_max_drawdown": ("买入持有最大回撤", "Buy And Hold Max Drawdown", "percent"),
}


DATA_QUALITY_LABELS = {
    "provider": ("数据源", "Provider"),
    "symbol": ("标的代码", "Symbol"),
    "market": ("市场", "Market"),
    "interval": ("周期", "Interval"),
    "request_time": ("请求时间", "Request Time"),
    "first_datetime": ("首条时间", "First Datetime"),
    "last_datetime": ("末条时间", "Last Datetime"),
    "row_count": ("数据行数", "Row Count"),
    "missing_values": ("缺失值数量", "Missing Values"),
    "duplicate_datetimes": ("重复时间戳数量", "Duplicate Datetimes"),
    "checksum": ("数据校验码", "Checksum"),
    "quality_status": ("质量状态", "Quality Status"),
    "notes": ("备注", "Notes"),
    "summary": ("摘要", "Summary"),
}




PARAMETER_STABILITY_LABELS = {
    "score_metric": ("评分指标", "Score Metric"),
    "best_score": ("最佳分数", "Best Score"),
    "top_quantile_mean": ("顶部区间均值", "Top Quantile Mean"),
    "top_quantile_threshold": ("顶部区间阈值", "Top Quantile Threshold"),
    "neighbor_mean": ("邻近参数均值", "Neighbor Mean"),
    "neighbor_count": ("邻近参数数量", "Neighbor Count"),
    "parameter_coverage": ("参数覆盖率", "Parameter Coverage"),
    "stability_status": ("稳定性状态", "Stability Status"),
    "notes": ("说明", "Notes"),
}


VALIDATION_LABELS = {
    "split_datetime": ("切分时间", "Split Datetime"),
    "train_rows": ("训练行数", "Train Rows"),
    "test_rows": ("测试行数", "Test Rows"),
    "score_metric": ("评分指标", "Score Metric"),
    "best_train_run_id": ("训练期最佳运行", "Best Train Run Id"),
    "best_params": ("最佳参数", "Best Parameters"),
    "train_score": ("训练分数", "Train Score"),
    "test_score": ("测试分数", "Test Score"),
    "train_total_return": ("训练总收益", "Train Total Return"),
    "test_total_return": ("测试总收益", "Test Total Return"),
    "train_max_drawdown": ("训练最大回撤", "Train Maximum Drawdown"),
    "test_max_drawdown": ("测试最大回撤", "Test Maximum Drawdown"),
    "generalization_ratio": ("泛化比率", "Generalization Ratio"),
    "validation_status": ("验证状态", "Validation Status"),
    "notes": ("说明", "Notes"),
}


WALK_FORWARD_LABELS = {
    "window_count": ("窗口数量", "Window Count"),
    "pass_count": ("通过数量", "Pass Count"),
    "watch_count": ("观察数量", "Watch Count"),
    "failed_count": ("失败数量", "Failed Count"),
    "average_train_score": ("平均训练分数", "Average Train Score"),
    "average_test_score": ("平均测试分数", "Average Test Score"),
    "average_generalization_ratio": ("平均泛化比率", "Average Generalization Ratio"),
    "validation_status": ("验证状态", "Validation Status"),
    "notes": ("说明", "Notes"),
}


TRACEABILITY_LABELS = {
    "summary_path": ("实验汇总文件", "Summary Path"),
    "runs_path": ("运行明细文件", "Runs Path"),
    "stability_path": ("稳定性文件", "Stability Path"),
    "validation_path": ("样本外验证文件", "Train-Test Validation Path"),
    "walk_forward_path": ("滚动验证文件", "Walk-Forward Validation Path"),
}

from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import sqlite3
from contextlib import contextmanager
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd

from pfi_os.config import DATA_DIR, REPORT_ROOT_DIR
from pfi_os.integrations.external_systems import collect_industry_reports
from pfi_os.integrations.holding_symbol_map import build_entity_registry, holdings_symbol_proxy_frame
from pfi_os.integrations.holdings_book import load_current_holdings
from pfi_os.research.validation_queue import VALIDATION_QUEUE_PATH, create_validation_task, load_validation_tasks
from pfi_os.storage import atomic_write_json, locked_json_update


RESEARCH_BUS_DIR = DATA_DIR / "researchBus"
RESEARCH_BUS_DB_PATH = RESEARCH_BUS_DIR / "ResearchBus.sqlite"
RESEARCH_BUS_SNAPSHOT_PATH = RESEARCH_BUS_DIR / "ResearchBusSnapshot.json"
DEFAULT_AI_RESEARCH_ROOT = Path("$PFI_AI_RESEARCH_ROOT")
RESEARCH_BUS_SCHEMA_VERSION = "ResearchBusV1"
REPORT_SUFFIXES = {".md", ".txt", ".docx", ".doc", ".pdf"}
VALIDATION_KEYWORDS = (
    "验证",
    "回测",
    "趋势",
    "动量",
    "均线",
    "rsi",
    "macd",
    "估值",
    "财报",
    "公告",
    "催化",
    "政策",
    "行业",
    "需求",
    "价格",
    "放量",
    "回撤",
    "失效",
    "观察",
    "thesis",
    "signal",
    "backtest",
    "risk",
)
US_TICKER_EXCLUDE = {
    "A",
    "AI",
    "API",
    "CEO",
    "CFO",
    "CPI",
    "ETF",
    "FOMC",
    "GDP",
    "HK",
    "IPO",
    "MACD",
    "NAV",
    "PE",
    "PEG",
    "PMI",
    "PRD",
    "QE",
    "QQQ",
    "RMB",
    "ROE",
    "RSI",
    "SOP",
    "USD",
    "US",
}


@dataclass(frozen=True)
class ResearchBusSyncResult:
    synced_at: str
    db_path: str
    registered_systems: int = 0
    system_artifacts: int = 0
    reports: int = 0
    validation_tasks: int = 0
    pfi_os_results: int = 0
    holdings: int = 0
    holding_symbol_mappings: int = 0
    portfolio_transactions: int = 0
    events: int = 0
    snapshot_path: str = ""
    ai_bridge_path: str = ""
    warnings: tuple[str, ...] = ()

    def to_dict(self) -> dict[str, Any]:
        return {
            "synced_at": self.synced_at,
            "db_path": self.db_path,
            "registered_systems": self.registered_systems,
            "system_artifacts": self.system_artifacts,
            "reports": self.reports,
            "validation_tasks": self.validation_tasks,
            "pfi_os_results": self.pfi_os_results,
            "holdings": self.holdings,
            "holding_symbol_mappings": self.holding_symbol_mappings,
            "portfolio_transactions": self.portfolio_transactions,
            "events": self.events,
            "snapshot_path": self.snapshot_path,
            "ai_bridge_path": self.ai_bridge_path,
            "warnings": list(self.warnings),
        }


def research_bus_db_path(path: Path | str | None = None) -> Path:
    if path is not None:
        return Path(path).expanduser()
    configured = os.getenv("PFI_RESEARCH_BUS_DB", "").strip()
    if configured:
        return Path(configured).expanduser()
    return RESEARCH_BUS_DB_PATH


def initialize_research_bus(path: Path | str | None = None) -> Path:
    db_path = research_bus_db_path(path)
    db_path.parent.mkdir(parents=True, exist_ok=True)
    with _connect(db_path) as conn:
        _create_schema(conn)
    return db_path


def sync_all_research_bus(
    report_root: Path | str | None = None,
    industry_report_root: Path | str | None = None,
    db_path: Path | str | None = None,
    *,
    push_validation_queue: bool = True,
    ai_research_root: Path | str | None = None,
) -> ResearchBusSyncResult:
    target_db = initialize_research_bus(db_path)
    try:
        from pfi_os.integrations.system_orchestrator import register_default_systems, sync_default_system_artifacts

        registry_result = register_default_systems(target_db)
        artifact_result = sync_default_system_artifacts(target_db)
    except Exception as exc:
        registry_result = {"registered_systems": 0, "warnings": [f"系统注册表同步失败：{exc}"]}
        artifact_result = {"system_artifacts": 0, "warnings": []}
    reports_result = sync_industry_reports_to_bus(industry_report_root, target_db, push_validation_queue=push_validation_queue)
    results_result = sync_pfi_os_results_to_bus(report_root, target_db, ai_research_root=ai_research_root)
    holdings_result = sync_holdings_to_bus(db_path=target_db)
    mappings_result = sync_holding_symbol_mappings_to_bus(db_path=target_db)
    transactions_result = sync_portfolio_transactions_to_bus(db_path=target_db)
    try:
        from pfi_os.integrations.consumer_behavior import sync_consumer_behavior_state

        consumer_result = sync_consumer_behavior_state(bus_db_path=target_db)
    except Exception as exc:
        consumer_result = ResearchBusSyncResult(synced_at=_now(), db_path=str(target_db), warnings=(f"消费行为系统同步失败：{exc}",))
    snapshot_path = export_research_bus_snapshot(target_db)
    warnings = [
        *registry_result.get("warnings", []),
        *artifact_result.get("warnings", []),
        *reports_result.warnings,
        *results_result.warnings,
        *holdings_result.warnings,
        *mappings_result.warnings,
        *transactions_result.warnings,
        *consumer_result.warnings,
    ]
    return ResearchBusSyncResult(
        synced_at=_now(),
        db_path=str(target_db),
        registered_systems=int(registry_result.get("registered_systems", 0)),
        system_artifacts=int(artifact_result.get("system_artifacts", 0)),
        reports=reports_result.reports,
        validation_tasks=reports_result.validation_tasks,
        pfi_os_results=results_result.pfi_os_results,
        holdings=holdings_result.holdings,
        holding_symbol_mappings=mappings_result.holding_symbol_mappings,
        portfolio_transactions=transactions_result.portfolio_transactions,
        events=reports_result.events
        + results_result.events
        + holdings_result.events
        + mappings_result.events
        + transactions_result.events
        + consumer_result.events,
        snapshot_path=str(snapshot_path),
        ai_bridge_path=results_result.ai_bridge_path,
        warnings=tuple(warnings),
    )


def sync_industry_reports_to_bus(
    report_root: Path | str | None = None,
    db_path: Path | str | None = None,
    *,
    push_validation_queue: bool = False,
) -> ResearchBusSyncResult:
    target_db = initialize_research_bus(db_path)
    reports = collect_industry_reports(report_root)
    report_count = 0
    task_count = 0
    warnings: list[str] = []
    with _connect(target_db) as conn:
        _upsert_system_state(
            conn,
            "AI-Research-System",
            "Ready" if not reports.empty else "ConfiguredNoData",
            str(Path(report_root).expanduser()) if report_root is not None else "",
            {"report_count": int(len(reports)), "purpose": "行研报告索引与验证任务来源"},
        )
        for row in reports.to_dict("records"):
            file_path = Path(str(row.get("path", ""))).expanduser()
            if not file_path.exists() or file_path.suffix.lower() not in REPORT_SUFFIXES:
                continue
            parsed = parse_report_file(file_path)
            report_id = _stable_id("report", str(file_path.resolve()), parsed["content_hash"])
            symbols = extract_symbols(parsed["text"], fallback_text=file_path.name)
            topics = extract_topics(parsed["text"], fallback_text=str(row.get("category", "")))
            _upsert_research_report(conn, report_id, row, parsed, symbols, topics)
            report_count += 1
            candidates = extract_validation_candidates(parsed["text"], row, report_id, file_path, symbols)
            if not candidates and file_path.suffix.lower() == ".pdf":
                warnings.append(f"PDF 正文未解析，仅完成索引：{file_path.name}")
            for candidate in candidates:
                _upsert_validation_task(conn, candidate)
                task_count += 1
        _record_event(
            conn,
            "industry_reports_sync",
            "AI-Research-System",
            "ResearchBus",
            "success",
            f"reports={report_count}; validation_tasks={task_count}",
        )
    pushed = 0
    if push_validation_queue:
        pushed = push_bus_validation_tasks_to_pfi_os_queue(target_db)
    return ResearchBusSyncResult(
        synced_at=_now(),
        db_path=str(target_db),
        reports=report_count,
        validation_tasks=task_count if not push_validation_queue else pushed,
        events=1,
        warnings=tuple(warnings),
    )


def sync_pfi_os_results_to_bus(
    report_root: Path | str | None = None,
    db_path: Path | str | None = None,
    *,
    ai_research_root: Path | str | None = None,
) -> ResearchBusSyncResult:
    from pfi_os.reports.catalog import run_metadata_summaries_frame

    target_db = initialize_research_bus(db_path)
    runs = run_metadata_summaries_frame(report_root or REPORT_ROOT_DIR)
    result_count = 0
    warnings: list[str] = []
    exported_payloads: list[dict[str, Any]] = []
    with _connect(target_db) as conn:
        _upsert_system_state(
            conn,
            "PFIOS",
            "Ready" if not runs.empty else "ConfiguredNoData",
            str(Path(report_root).expanduser()) if report_root is not None else str(REPORT_ROOT_DIR),
            {"run_metadata_count": int(len(runs)), "purpose": "回测、研究门禁与报告结果来源"},
        )
        for row in runs.to_dict("records"):
            metadata_path = Path(str(row.get("metadata_path", ""))).expanduser()
            payload = _safe_json_file(metadata_path)
            result_id = _stable_id("pfi_osResult", str(metadata_path.resolve()) if metadata_path.exists() else json.dumps(row, sort_keys=True, ensure_ascii=False))
            _upsert_pfi_os_result(conn, result_id, row, payload)
            exported_payloads.append({"result_id": result_id, **row, "payload": payload})
            result_count += 1
        _record_event(conn, "pfi_os_results_sync", "PFIOS", "ResearchBus", "success", f"results={result_count}")
    bridge_path = ""
    if ai_research_root is not None:
        bridge_path = str(write_pfi_os_results_for_ai_research(exported_payloads, ai_research_root))
    elif DEFAULT_AI_RESEARCH_ROOT.exists():
        bridge_path = str(write_pfi_os_results_for_ai_research(exported_payloads, DEFAULT_AI_RESEARCH_ROOT))
    else:
        warnings.append("未找到行研系统目录，已只写入共享研究数据总线。")
    return ResearchBusSyncResult(
        synced_at=_now(),
        db_path=str(target_db),
        pfi_os_results=result_count,
        events=1,
        ai_bridge_path=bridge_path,
        warnings=tuple(warnings),
    )


def sync_holdings_to_bus(
    holdings: pd.DataFrame | None = None,
    db_path: Path | str | None = None,
) -> ResearchBusSyncResult:
    target_db = initialize_research_bus(db_path)
    frame = holdings if isinstance(holdings, pd.DataFrame) else load_current_holdings()
    holding_count = 0
    with _connect(target_db) as conn:
        _upsert_system_state(
            conn,
            "HoldingsMaster",
            "Ready" if not frame.empty else "ConfiguredNoData",
            "",
            {"holding_count": int(len(frame)), "purpose": "跨系统统一持仓主数据"},
        )
        if not frame.empty:
            source_systems = sorted({str(item).strip() for item in frame.get("source_system", pd.Series(dtype=str)).dropna() if str(item).strip()})
            for source_system in source_systems:
                conn.execute("DELETE FROM holdings_master WHERE source_system=?", (source_system,))
            for row in frame.to_dict("records"):
                _upsert_holding(conn, row)
                holding_count += 1
        _record_event(conn, "holdings_sync", "PFIOS", "ResearchBus", "success", f"holdings={holding_count}")
    return ResearchBusSyncResult(synced_at=_now(), db_path=str(target_db), holdings=holding_count, events=1)


def sync_holding_symbol_mappings_to_bus(
    holdings: pd.DataFrame | None = None,
    db_path: Path | str | None = None,
) -> ResearchBusSyncResult:
    target_db = initialize_research_bus(db_path)
    frame = holdings if isinstance(holdings, pd.DataFrame) else load_current_holdings()
    mappings = holdings_symbol_proxy_frame(frame)
    entity_registry = build_entity_registry(frame)
    mapping_count = 0
    with _connect(target_db) as conn:
        _upsert_system_state(
            conn,
            "HoldingSymbolMappings",
            "Ready" if not mappings.empty else "ConfiguredNoData",
            "",
            {
                "mapping_count": int(len(mappings)),
                "mapped_count": int((mappings.get("status", pd.Series(dtype=str)) != "MissingSymbol").sum()) if not mappings.empty else 0,
                "entity_registry_schema": entity_registry["schema"],
                "entity_status_counts": entity_registry["status_counts"],
                "entity_market_counts": entity_registry["market_counts"],
                "purpose": "持仓名称到行情代码或 ETF/指数代理的跨系统映射；仅用于研究观察和情绪分析。",
            },
        )
        conn.execute("DELETE FROM holding_symbol_mappings WHERE source_system='PFIOS'")
        for row in mappings.to_dict("records"):
            _upsert_holding_symbol_mapping(conn, row)
            mapping_count += 1
        _record_event(
            conn,
            "holding_symbol_mappings_sync",
            "PFIOS",
            "ResearchBus",
            "success",
            f"holding_symbol_mappings={mapping_count}",
        )
    return ResearchBusSyncResult(
        synced_at=_now(),
        db_path=str(target_db),
        holding_symbol_mappings=mapping_count,
        events=1,
    )


def sync_portfolio_transactions_to_bus(
    transactions: pd.DataFrame | None = None,
    db_path: Path | str | None = None,
) -> ResearchBusSyncResult:
    target_db = initialize_research_bus(db_path)
    frame = transactions if isinstance(transactions, pd.DataFrame) else load_default_portfolio_transactions()
    transaction_count = 0
    with _connect(target_db) as conn:
        _upsert_system_state(
            conn,
            "PortfolioTransactions",
            "Ready" if not frame.empty else "ConfiguredNoData",
            "",
            {"transaction_count": int(len(frame)), "purpose": "跨系统统一交易记录与待确认订单"},
        )
        if not frame.empty:
            for row in frame.to_dict("records"):
                _upsert_portfolio_transaction(conn, row)
                transaction_count += 1
        _record_event(
            conn,
            "portfolio_transactions_sync",
            "PFIOS",
            "ResearchBus",
            "success",
            f"portfolio_transactions={transaction_count}",
        )
    return ResearchBusSyncResult(
        synced_at=_now(),
        db_path=str(target_db),
        portfolio_transactions=transaction_count,
        events=1,
    )


def load_default_portfolio_transactions(paths: tuple[Path | str, ...] | None = None) -> pd.DataFrame:
    roots = tuple(Path(item).expanduser() for item in paths) if paths is not None else _default_transaction_roots()
    rows: list[dict[str, Any]] = []
    seen_files: set[str] = set()
    for root in roots:
        files = _transaction_source_files(root)
        for file_path in files:
            key = str(file_path.resolve())
            if key in seen_files:
                continue
            seen_files.add(key)
            try:
                frame = pd.read_csv(file_path)
            except Exception:
                continue
            for row in frame.to_dict("records"):
                rows.append(_normalize_transaction_row(row, file_path))
    if not rows:
        return pd.DataFrame(columns=_transaction_columns())
    data = pd.DataFrame(rows)
    data = data[_transaction_columns()].copy()
    data["order_amount"] = pd.to_numeric(data["order_amount"], errors="coerce").fillna(0.0)
    data["confirmed_amount"] = pd.to_numeric(data["confirmed_amount"], errors="coerce").fillna(0.0)
    data["fee"] = pd.to_numeric(data["fee"], errors="coerce").fillna(0.0)
    data["_key"] = data.apply(
        lambda row: "|".join(
            str(row.get(part, "")).strip()
            for part in ["trade_date", "order_time", "name", "side", "order_amount"]
        ),
        axis=1,
    )
    data["_quality_priority"] = data["quality_status"].map(_transaction_quality_priority).fillna(90)
    data["_source_priority"] = data["source_system"].map(_transaction_source_priority).fillna(90)
    data["_updated_ts"] = pd.to_datetime(data["updated_at"], errors="coerce")
    data = data.sort_values(["_key", "_quality_priority", "_source_priority", "_updated_ts"], ascending=[True, True, True, False])
    data = data.groupby("_key", as_index=False).head(1).drop(columns=["_key", "_quality_priority", "_source_priority", "_updated_ts"])
    return data.sort_values(["trade_date", "order_time", "name"], ascending=[False, False, True]).reset_index(drop=True)


def push_bus_validation_tasks_to_pfi_os_queue(
    db_path: Path | str | None = None,
    queue_path: Path | str | None = None,
) -> int:
    target_db = initialize_research_bus(db_path)
    target_queue_path = Path(queue_path).expanduser() if queue_path is not None else VALIDATION_QUEUE_PATH
    with _connect(target_db) as conn:
        rows = conn.execute(
            """
            SELECT task_id, source_report_path, source_paragraph, research_topic, symbol, market,
                   signal_to_validate, sample_period, cost_assumption, benchmark, status, validation_report_path
            FROM validation_tasks
            ORDER BY updated_at DESC, created_at DESC
            """
        ).fetchall()
    existing = load_validation_tasks(target_queue_path)
    existing_keys = {_validation_queue_key(task.to_dict()) for task in existing}
    pending_records: list[dict[str, Any]] = []
    for row in rows:
        payload = {
            "task_id": row["task_id"],
            "source_report": row["source_report_path"],
            "source_paragraph": row["source_paragraph"],
            "research_topic": row["research_topic"],
            "symbol": row["symbol"],
            "market": row["market"],
            "signal_to_validate": row["signal_to_validate"],
            "sample_period": row["sample_period"],
            "cost_assumption": row["cost_assumption"],
            "benchmark": row["benchmark"],
            "status": _normalize_queue_status(row["status"]),
            "validation_report_path": row["validation_report_path"],
            "notes": "由统一研究数据总线自动生成；正式使用前仍需人工确认研究假设和成本参数。",
        }
        key = _validation_queue_key(payload)
        if key in existing_keys:
            continue
        pending_records.append(create_validation_task(payload).to_dict())
        existing_keys.add(key)
    if not pending_records:
        return 0
    added_count = 0

    def append_tasks(records: list[dict[str, Any]]) -> list[dict[str, Any]]:
        nonlocal added_count
        clean_records = [item for item in records if isinstance(item, dict)]
        keys = {_validation_queue_key(item) for item in clean_records}
        additions = []
        for item in pending_records:
            key = _validation_queue_key(item)
            if key in keys:
                continue
            additions.append(item)
            keys.add(key)
        added_count = len(additions)
        return clean_records + additions

    locked_json_update(target_queue_path, [], append_tasks, expected_type=list)
    return added_count


def parse_report_file(path: Path | str) -> dict[str, str]:
    file_path = Path(path).expanduser()
    suffix = file_path.suffix.lower()
    text = ""
    warning = ""
    if suffix in {".md", ".txt"}:
        text = file_path.read_text(encoding="utf-8", errors="ignore")
    elif suffix == ".docx":
        text = _read_docx_text(file_path)
    elif suffix == ".pdf":
        text = _read_pdf_text(file_path)
        if not text:
            warning = _pdf_parse_warning()
    elif suffix == ".doc":
        warning = "Legacy .doc body parsing is not enabled; indexed by filename and metadata only."
    content_hash = _hash_file(file_path) if file_path.exists() else ""
    return {
        "text": _compact_whitespace(text),
        "content_hash": content_hash,
        "parser_warning": warning,
        "summary": _summary_text(text, file_path.name),
    }


def extract_symbols(text: str, fallback_text: str = "") -> list[str]:
    content = f"{text}\n{fallback_text}"
    symbols: set[str] = set()
    for match in re.findall(r"\b(?:SH|SZ)\.?\d{6}\b|\b\d{6}\.(?:SH|SZ|SS)\b|\b[036]\d{5}\b", content, flags=re.IGNORECASE):
        symbols.add(_normalize_symbol(match))
    for match in re.findall(r"\b(?:HK)\.?\d{4,5}\b|\b\d{4,5}\.HK\b", content, flags=re.IGNORECASE):
        symbols.add(_normalize_symbol(match))
    for match in re.findall(r"\b[A-Z]{1,5}\b", content):
        if match in US_TICKER_EXCLUDE or match in {"SH", "SZ", "SS", "HK"}:
            continue
        if len(match) == 1 and not re.search(rf"(NASDAQ|NYSE|AMEX|美股|US)[^\n]{{0,20}}\b{match}\b", content, flags=re.IGNORECASE):
            continue
        symbols.add(match)
    return sorted(symbols)


def extract_topics(text: str, fallback_text: str = "") -> list[str]:
    content = f"{text}\n{fallback_text}"
    topics: set[str] = set()
    for keyword in ["人工智能", "半导体", "银行", "黄金", "能源", "消费", "医药", "互联网", "地产", "新能源", "港股", "美股", "A股", "政策", "财报", "估值"]:
        if keyword.lower() in content.lower():
            topics.add(keyword)
    return sorted(topics)


def extract_validation_candidates(
    text: str,
    report_row: dict[str, Any],
    report_id: str,
    path: Path,
    symbols: list[str] | None = None,
) -> list[dict[str, Any]]:
    symbols = symbols or extract_symbols(text, fallback_text=path.name)
    paragraphs = _candidate_paragraphs(text)
    if not paragraphs and symbols:
        paragraphs = [f"{path.name} 中提到 {', '.join(symbols[:6])}，需由 PFIOS 补充验证。"]
    candidates: list[dict[str, Any]] = []
    for paragraph in paragraphs[:60]:
        paragraph_symbols = extract_symbols(paragraph, fallback_text="")
        active_symbols = paragraph_symbols or symbols[:8] or [""]
        for symbol in active_symbols[:8]:
            market = infer_market(symbol)
            task_id = _stable_id("validationTask", report_id, symbol, paragraph)
            candidates.append(
                {
                    "task_id": task_id,
                    "source_system": "AI-Research-System",
                    "source_report_id": report_id,
                    "source_report_path": str(path),
                    "source_paragraph": _truncate(paragraph, 800),
                    "research_topic": _research_topic(paragraph, report_row),
                    "symbol": symbol,
                    "market": market,
                    "signal_to_validate": _signal_to_validate(paragraph),
                    "sample_period": "由 PFIOS 回测页面选择；建议至少覆盖一轮牛熊或完整行业周期。",
                    "cost_assumption": "使用 PFIOS 当前手续费、滑点、冲击成本设置；报告中必须披露公式和假设。",
                    "benchmark": benchmark_for_market(market),
                    "status": "待验证",
                    "validation_report_path": "",
                    "created_at": _now(),
                    "updated_at": _now(),
                }
            )
    return candidates


def infer_market(symbol: str) -> str:
    text = str(symbol or "").upper().strip()
    if not text:
        return ""
    if text.endswith(".HK") or text.startswith("HK."):
        return "HK"
    if text.endswith(".SH") or text.endswith(".SZ") or text.endswith(".SS") or re.fullmatch(r"[036]\d{5}", text):
        return "CN"
    return "US"


def benchmark_for_market(market: str) -> str:
    return {"CN": "沪深300", "HK": "恒生指数", "US": "标普500"}.get(str(market).upper(), "自定义基准")


def research_bus_status_frame(db_path: Path | str | None = None) -> pd.DataFrame:
    target_db = initialize_research_bus(db_path)
    with _connect(target_db) as conn:
        rows = conn.execute(
            """
            SELECT system_name, status, root_path, last_sync_at, summary_json
            FROM system_state
            ORDER BY last_sync_at DESC
            """
        ).fetchall()
        counts = {
            "system_registry": _table_count(conn, "system_registry"),
            "system_artifacts": _table_count(conn, "system_artifacts"),
            "orchestration_runs": _table_count(conn, "orchestration_runs"),
            "research_reports": _table_count(conn, "research_reports"),
            "validation_tasks": _table_count(conn, "validation_tasks"),
            "pfi_os_results": _table_count(conn, "pfi_os_results"),
            "holdings_master": _table_count(conn, "holdings_master"),
            "holding_symbol_mappings": _table_count(conn, "holding_symbol_mappings"),
            "portfolio_transactions": _table_count(conn, "portfolio_transactions"),
            "holding_update_candidates": _table_count(conn, "holding_update_candidates"),
            "consumer_behavior_state": _table_count(conn, "consumer_behavior_state"),
            "independent_validation_runs": _table_count(conn, "independent_validation_runs"),
            "independent_validation_shards": _table_count(conn, "independent_validation_shards"),
            "bus_api_requests": _table_count(conn, "bus_api_requests"),
            "bus_chat_inputs": _table_count(conn, "bus_chat_inputs"),
            "bus_system_outbox": _table_count(conn, "bus_system_outbox"),
            "bus_heartbeats": _table_count(conn, "bus_heartbeats"),
            "sync_events": _table_count(conn, "sync_events"),
        }
    records = [dict(row) for row in rows]
    records.append(
        {
            "system_name": "ResearchBus",
            "status": "Ready",
            "root_path": str(target_db),
            "last_sync_at": _now(),
            "summary_json": json.dumps(counts, ensure_ascii=False, sort_keys=True),
        }
    )
    return pd.DataFrame(records)


def bus_validation_task_frame(db_path: Path | str | None = None) -> pd.DataFrame:
    target_db = initialize_research_bus(db_path)
    with _connect(target_db) as conn:
        rows = conn.execute(
            """
            SELECT created_at, research_topic, symbol, market, status, validation_report_path,
                   source_report_path, source_paragraph, signal_to_validate
            FROM validation_tasks
            ORDER BY updated_at DESC, created_at DESC
            """
        ).fetchall()
    return pd.DataFrame([dict(row) for row in rows])


def bus_pfi_os_results_frame(db_path: Path | str | None = None) -> pd.DataFrame:
    target_db = initialize_research_bus(db_path)
    with _connect(target_db) as conn:
        rows = conn.execute(
            """
            SELECT created_at, strategy_id, symbol, market, total_return, annualized_return,
                   max_drawdown, sharpe, research_status, decision_quality_score, report_path, metadata_path
            FROM pfi_os_results
            ORDER BY updated_at DESC, created_at DESC
            """
        ).fetchall()
    return pd.DataFrame([dict(row) for row in rows])


def portfolio_transactions_frame(db_path: Path | str | None = None) -> pd.DataFrame:
    target_db = initialize_research_bus(db_path)
    with _connect(target_db) as conn:
        rows = conn.execute(
            """
            SELECT trade_date, order_time, timezone, source_system, account, symbol, name, market,
                   side, order_type, order_amount, confirmed_amount, fee, status, quality_status,
                   source_path, evidence_frame, updated_at
            FROM portfolio_transactions
            ORDER BY trade_date DESC, order_time DESC, updated_at DESC
            """
        ).fetchall()
    return pd.DataFrame([dict(row) for row in rows])


def holding_symbol_mappings_frame(db_path: Path | str | None = None) -> pd.DataFrame:
    target_db = initialize_research_bus(db_path)
    with _connect(target_db) as conn:
        rows = conn.execute(
            """
            SELECT holding_name, holding_market, original_symbol, proxy_symbol, proxy_name,
                   proxy_market, status, confidence, reason, source, updated_at
            FROM holding_symbol_mappings
            ORDER BY
                CASE status
                    WHEN 'ConfirmedSymbol' THEN 1
                    WHEN 'ProxyMapped' THEN 2
                    ELSE 3
                END,
                holding_name ASC
            """
        ).fetchall()
    return pd.DataFrame([dict(row) for row in rows])


def holding_update_candidates_frame(db_path: Path | str | None = None, limit: int = 500) -> pd.DataFrame:
    target_db = initialize_research_bus(db_path)
    with _connect(target_db) as conn:
        rows = conn.execute(
            """
            SELECT candidate_id, source_system, account, candidate_type, status, quality_status,
                   content_text, attachments_json, extracted_symbols_json, source_request_id,
                   payload_json, created_at, updated_at
            FROM holding_update_candidates
            ORDER BY updated_at DESC, created_at DESC
            LIMIT ?
            """,
            (int(limit),),
        ).fetchall()
    frame = pd.DataFrame([dict(row) for row in rows])
    if frame.empty or "payload_json" not in frame.columns:
        return frame
    parser_statuses: list[str] = []
    holding_counts: list[int] = []
    transaction_counts: list[int] = []
    for value in frame["payload_json"]:
        payload = _safe_json_text(value)
        reports = payload.get("attachment_parser_reports", []) if isinstance(payload, dict) else []
        parser_statuses.append(",".join(str(report.get("status", "")) for report in reports if isinstance(report, dict)))
        holding_counts.append(len(payload.get("holdings", [])) if isinstance(payload.get("holdings", []), list) else 0)
        transaction_counts.append(len(payload.get("portfolio_transactions", [])) if isinstance(payload.get("portfolio_transactions", []), list) else 0)
    frame["parser_status"] = parser_statuses
    frame["structured_holding_count"] = holding_counts
    frame["structured_transaction_count"] = transaction_counts
    return frame


def export_research_bus_snapshot(db_path: Path | str | None = None, path: Path | str | None = None) -> Path:
    target_db = initialize_research_bus(db_path)
    snapshot_path = Path(path).expanduser() if path is not None else RESEARCH_BUS_SNAPSHOT_PATH
    payload: dict[str, Any] = {
        "schema": RESEARCH_BUS_SCHEMA_VERSION,
        "exported_at": _now(),
        "db_path": str(target_db),
        "tables": {},
    }
    with _connect(target_db) as conn:
        for table in [
            "system_state",
            "system_registry",
            "system_artifacts",
            "orchestration_runs",
            "research_reports",
            "validation_tasks",
            "pfi_os_results",
            "holdings_master",
            "holding_symbol_mappings",
            "portfolio_transactions",
            "holding_update_candidates",
            "consumer_behavior_state",
            "independent_validation_runs",
            "independent_validation_shards",
            "bus_api_requests",
            "bus_chat_inputs",
            "bus_system_outbox",
            "bus_heartbeats",
            "sync_events",
        ]:
            rows = conn.execute(f"SELECT * FROM {table} ORDER BY rowid DESC LIMIT 500").fetchall()
            payload["tables"][table] = [dict(row) for row in rows]
    atomic_write_json(snapshot_path, payload)
    return snapshot_path


def write_pfi_os_results_for_ai_research(
    results: list[dict[str, Any]],
    ai_research_root: Path | str,
) -> Path:
    root = Path(ai_research_root).expanduser()
    out_path = root / "data" / "report_artifacts" / "pfi_os_bridge" / "PFIOSResults.json"
    payload = {
        "schema": RESEARCH_BUS_SCHEMA_VERSION,
        "exported_at": _now(),
        "source_system": "PFIOS",
        "purpose": "供行研系统读取 PFIOS 回测、门禁、决策质量和报告路径。",
        "results": results[-500:],
    }
    atomic_write_json(out_path, payload)
    return out_path


@contextmanager
def _connect(path: Path | str):
    conn = sqlite3.connect(str(path), timeout=30)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA busy_timeout=30000")
    conn.execute("PRAGMA foreign_keys=ON")
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


def _create_schema(conn: sqlite3.Connection) -> None:
    conn.executescript(
        """
        CREATE TABLE IF NOT EXISTS system_state (
            system_name TEXT PRIMARY KEY,
            status TEXT NOT NULL,
            root_path TEXT NOT NULL DEFAULT '',
            last_sync_at TEXT NOT NULL,
            summary_json TEXT NOT NULL DEFAULT '{}'
        );

        CREATE TABLE IF NOT EXISTS system_registry (
            system_name TEXT PRIMARY KEY,
            role TEXT NOT NULL DEFAULT '',
            root_path TEXT NOT NULL DEFAULT '',
            standalone_command_json TEXT NOT NULL DEFAULT '[]',
            health_command_json TEXT NOT NULL DEFAULT '[]',
            sync_command_json TEXT NOT NULL DEFAULT '[]',
            capabilities_json TEXT NOT NULL DEFAULT '[]',
            outputs_json TEXT NOT NULL DEFAULT '[]',
            status TEXT NOT NULL DEFAULT '',
            last_seen_at TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS system_artifacts (
            artifact_id TEXT PRIMARY KEY,
            system_name TEXT NOT NULL DEFAULT '',
            artifact_type TEXT NOT NULL DEFAULT '',
            title TEXT NOT NULL DEFAULT '',
            path TEXT NOT NULL DEFAULT '',
            content_hash TEXT NOT NULL DEFAULT '',
            size_bytes INTEGER NOT NULL DEFAULT 0,
            modified_at TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS orchestration_runs (
            run_id TEXT PRIMARY KEY,
            parent_system TEXT NOT NULL DEFAULT 'PFIOS',
            target_system TEXT NOT NULL DEFAULT '',
            action TEXT NOT NULL DEFAULT '',
            mode TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT '',
            started_at TEXT NOT NULL DEFAULT '',
            completed_at TEXT NOT NULL DEFAULT '',
            command_json TEXT NOT NULL DEFAULT '[]',
            exit_code INTEGER NOT NULL DEFAULT 0,
            stdout_path TEXT NOT NULL DEFAULT '',
            stderr_path TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS research_reports (
            report_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL,
            report_type TEXT NOT NULL DEFAULT '',
            title TEXT NOT NULL,
            report_date TEXT NOT NULL DEFAULT '',
            period TEXT NOT NULL DEFAULT '',
            path TEXT NOT NULL,
            content_hash TEXT NOT NULL DEFAULT '',
            summary TEXT NOT NULL DEFAULT '',
            symbols_json TEXT NOT NULL DEFAULT '[]',
            topics_json TEXT NOT NULL DEFAULT '[]',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS validation_tasks (
            task_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL,
            source_report_id TEXT NOT NULL DEFAULT '',
            source_report_path TEXT NOT NULL DEFAULT '',
            source_paragraph TEXT NOT NULL DEFAULT '',
            research_topic TEXT NOT NULL DEFAULT '',
            symbol TEXT NOT NULL DEFAULT '',
            market TEXT NOT NULL DEFAULT '',
            signal_to_validate TEXT NOT NULL DEFAULT '',
            sample_period TEXT NOT NULL DEFAULT '',
            cost_assumption TEXT NOT NULL DEFAULT '',
            benchmark TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT '待验证',
            validation_report_path TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS pfi_os_results (
            result_id TEXT PRIMARY KEY,
            report_path TEXT NOT NULL DEFAULT '',
            metadata_path TEXT NOT NULL DEFAULT '',
            strategy_id TEXT NOT NULL DEFAULT '',
            symbol TEXT NOT NULL DEFAULT '',
            market TEXT NOT NULL DEFAULT '',
            total_return REAL NOT NULL DEFAULT 0,
            annualized_return REAL NOT NULL DEFAULT 0,
            max_drawdown REAL NOT NULL DEFAULT 0,
            sharpe REAL NOT NULL DEFAULT 0,
            research_status TEXT NOT NULL DEFAULT '',
            decision_quality_score REAL NOT NULL DEFAULT 0,
            data_quality_status TEXT NOT NULL DEFAULT '',
            cross_validation_status TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            payload_json TEXT NOT NULL DEFAULT '{}'
        );

        CREATE TABLE IF NOT EXISTS holdings_master (
            holding_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT '',
            account TEXT NOT NULL DEFAULT '',
            symbol TEXT NOT NULL DEFAULT '',
            name TEXT NOT NULL DEFAULT '',
            market TEXT NOT NULL DEFAULT '',
            asset_type TEXT NOT NULL DEFAULT '',
            quantity REAL NOT NULL DEFAULT 0,
            cost_basis REAL NOT NULL DEFAULT 0,
            position_value REAL NOT NULL DEFAULT 0,
            unrealized_pnl REAL NOT NULL DEFAULT 0,
            weight REAL NOT NULL DEFAULT 0,
            as_of TEXT NOT NULL DEFAULT '',
            source_path TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS holding_symbol_mappings (
            mapping_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT 'PFIOS',
            holding_name TEXT NOT NULL DEFAULT '',
            holding_market TEXT NOT NULL DEFAULT '',
            original_symbol TEXT NOT NULL DEFAULT '',
            proxy_symbol TEXT NOT NULL DEFAULT '',
            proxy_name TEXT NOT NULL DEFAULT '',
            proxy_market TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT '',
            confidence TEXT NOT NULL DEFAULT '',
            reason TEXT NOT NULL DEFAULT '',
            source TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS portfolio_transactions (
            transaction_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT '',
            account TEXT NOT NULL DEFAULT '',
            trade_date TEXT NOT NULL DEFAULT '',
            order_time TEXT NOT NULL DEFAULT '',
            timezone TEXT NOT NULL DEFAULT '',
            symbol TEXT NOT NULL DEFAULT '',
            name TEXT NOT NULL DEFAULT '',
            market TEXT NOT NULL DEFAULT '',
            asset_type TEXT NOT NULL DEFAULT '',
            side TEXT NOT NULL DEFAULT '',
            order_type TEXT NOT NULL DEFAULT '',
            order_amount REAL NOT NULL DEFAULT 0,
            confirmed_amount REAL NOT NULL DEFAULT 0,
            confirmed_units REAL NOT NULL DEFAULT 0,
            confirmed_nav REAL NOT NULL DEFAULT 0,
            fee REAL NOT NULL DEFAULT 0,
            status TEXT NOT NULL DEFAULT '',
            quality_status TEXT NOT NULL DEFAULT '',
            source_path TEXT NOT NULL DEFAULT '',
            evidence_frame TEXT NOT NULL DEFAULT '',
            notes TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS holding_update_candidates (
            candidate_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT '',
            account TEXT NOT NULL DEFAULT '',
            candidate_type TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'PendingReview',
            quality_status TEXT NOT NULL DEFAULT 'Candidate',
            content_text TEXT NOT NULL DEFAULT '',
            attachments_json TEXT NOT NULL DEFAULT '[]',
            extracted_symbols_json TEXT NOT NULL DEFAULT '[]',
            source_request_id TEXT NOT NULL DEFAULT '',
            source_chat_input_id TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS consumer_behavior_state (
            state_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT 'ConsumptionAnalysisSystem',
            db_path TEXT NOT NULL DEFAULT '',
            run_count INTEGER NOT NULL DEFAULT 0,
            transaction_count INTEGER NOT NULL DEFAULT 0,
            ledger_count INTEGER NOT NULL DEFAULT 0,
            latest_run_id TEXT NOT NULL DEFAULT '',
            latest_generated_at TEXT NOT NULL DEFAULT '',
            total_amount REAL NOT NULL DEFAULT 0,
            manual_review_count INTEGER NOT NULL DEFAULT 0,
            summary_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS sync_events (
            event_id TEXT PRIMARY KEY,
            event_type TEXT NOT NULL,
            source_system TEXT NOT NULL,
            target_system TEXT NOT NULL,
            status TEXT NOT NULL,
            message TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            created_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS independent_validation_runs (
            run_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT 'IndependentValidation',
            status TEXT NOT NULL DEFAULT '',
            mode TEXT NOT NULL DEFAULT '',
            manifest_path TEXT NOT NULL DEFAULT '',
            total_rows INTEGER NOT NULL DEFAULT 0,
            shard_count INTEGER NOT NULL DEFAULT 0,
            started_at TEXT NOT NULL DEFAULT '',
            completed_at TEXT NOT NULL DEFAULT '',
            output_path TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL
        );

        CREATE TABLE IF NOT EXISTS independent_validation_shards (
            shard_id TEXT PRIMARY KEY,
            run_id TEXT NOT NULL,
            shard_index INTEGER NOT NULL DEFAULT 0,
            source_path TEXT NOT NULL DEFAULT '',
            start_row INTEGER NOT NULL DEFAULT 0,
            end_row INTEGER NOT NULL DEFAULT 0,
            status TEXT NOT NULL DEFAULT '',
            payload_json TEXT NOT NULL DEFAULT '{}',
            updated_at TEXT NOT NULL,
            FOREIGN KEY(run_id) REFERENCES independent_validation_runs(run_id)
        );

        CREATE TABLE IF NOT EXISTS bus_api_requests (
            request_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT '',
            target_system TEXT NOT NULL DEFAULT '',
            request_type TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'Pending',
            priority INTEGER NOT NULL DEFAULT 5,
            payload_json TEXT NOT NULL DEFAULT '{}',
            response_json TEXT NOT NULL DEFAULT '{}',
            error_message TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL,
            updated_at TEXT NOT NULL,
            processed_at TEXT NOT NULL DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS bus_chat_inputs (
            input_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT '',
            author TEXT NOT NULL DEFAULT '',
            channel TEXT NOT NULL DEFAULT '',
            content_text TEXT NOT NULL DEFAULT '',
            attachments_json TEXT NOT NULL DEFAULT '[]',
            classification TEXT NOT NULL DEFAULT '',
            linked_request_id TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'Pending',
            payload_json TEXT NOT NULL DEFAULT '{}',
            created_at TEXT NOT NULL,
            processed_at TEXT NOT NULL DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS bus_system_outbox (
            message_id TEXT PRIMARY KEY,
            source_system TEXT NOT NULL DEFAULT '',
            target_system TEXT NOT NULL DEFAULT '',
            message_type TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT 'Pending',
            payload_json TEXT NOT NULL DEFAULT '{}',
            created_at TEXT NOT NULL,
            delivered_at TEXT NOT NULL DEFAULT ''
        );

        CREATE TABLE IF NOT EXISTS bus_heartbeats (
            system_name TEXT PRIMARY KEY,
            status TEXT NOT NULL DEFAULT '',
            capabilities_json TEXT NOT NULL DEFAULT '[]',
            payload_json TEXT NOT NULL DEFAULT '{}',
            last_seen_at TEXT NOT NULL
        );

        CREATE INDEX IF NOT EXISTS idx_validation_tasks_symbol ON validation_tasks(symbol, market);
        CREATE INDEX IF NOT EXISTS idx_system_registry_role ON system_registry(role, updated_at);
        CREATE INDEX IF NOT EXISTS idx_system_artifacts_system ON system_artifacts(system_name, modified_at);
        CREATE INDEX IF NOT EXISTS idx_orchestration_runs_target ON orchestration_runs(target_system, status, updated_at);
        CREATE INDEX IF NOT EXISTS idx_reports_date ON research_reports(report_date);
        CREATE INDEX IF NOT EXISTS idx_results_strategy ON pfi_os_results(strategy_id);
        CREATE INDEX IF NOT EXISTS idx_holdings_symbol ON holdings_master(symbol, market);
        CREATE INDEX IF NOT EXISTS idx_holding_symbol_mappings_proxy ON holding_symbol_mappings(proxy_symbol, proxy_market);
        CREATE INDEX IF NOT EXISTS idx_holding_symbol_mappings_name ON holding_symbol_mappings(holding_name, status);
        CREATE INDEX IF NOT EXISTS idx_portfolio_transactions_date ON portfolio_transactions(trade_date, order_time);
        CREATE INDEX IF NOT EXISTS idx_portfolio_transactions_name ON portfolio_transactions(name, side);
        CREATE INDEX IF NOT EXISTS idx_holding_update_candidates_status ON holding_update_candidates(status, updated_at);
        CREATE INDEX IF NOT EXISTS idx_consumer_behavior_state_updated ON consumer_behavior_state(updated_at);
        CREATE INDEX IF NOT EXISTS idx_independent_validation_runs_status ON independent_validation_runs(status, updated_at);
        CREATE INDEX IF NOT EXISTS idx_independent_validation_shards_run ON independent_validation_shards(run_id, shard_index);
        CREATE INDEX IF NOT EXISTS idx_bus_api_requests_target ON bus_api_requests(target_system, status, priority, created_at);
        CREATE INDEX IF NOT EXISTS idx_bus_chat_inputs_status ON bus_chat_inputs(status, created_at);
        CREATE INDEX IF NOT EXISTS idx_bus_system_outbox_target ON bus_system_outbox(target_system, status, created_at);
        CREATE INDEX IF NOT EXISTS idx_sync_events_created_at ON sync_events(created_at);
        """
    )


def _upsert_system_state(conn: sqlite3.Connection, system_name: str, status: str, root_path: str, summary: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT INTO system_state(system_name, status, root_path, last_sync_at, summary_json)
        VALUES (?, ?, ?, ?, ?)
        ON CONFLICT(system_name) DO UPDATE SET
            status=excluded.status,
            root_path=excluded.root_path,
            last_sync_at=excluded.last_sync_at,
            summary_json=excluded.summary_json
        """,
        (system_name, status, root_path, _now(), _json_dumps(summary)),
    )


def _upsert_research_report(
    conn: sqlite3.Connection,
    report_id: str,
    row: dict[str, Any],
    parsed: dict[str, str],
    symbols: list[str],
    topics: list[str],
) -> None:
    now = _now()
    conn.execute(
        """
        INSERT INTO research_reports(
            report_id, source_system, report_type, title, report_date, period, path, content_hash,
            summary, symbols_json, topics_json, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(report_id) DO UPDATE SET
            report_type=excluded.report_type,
            title=excluded.title,
            report_date=excluded.report_date,
            period=excluded.period,
            path=excluded.path,
            content_hash=excluded.content_hash,
            summary=excluded.summary,
            symbols_json=excluded.symbols_json,
            topics_json=excluded.topics_json,
            updated_at=excluded.updated_at
        """,
        (
            report_id,
            "AI-Research-System",
            str(row.get("category", "")),
            str(row.get("name", "")),
            str(row.get("report_date", "")),
            str(row.get("period", "")),
            str(row.get("path", "")),
            parsed.get("content_hash", ""),
            parsed.get("summary", ""),
            _json_dumps(symbols),
            _json_dumps(topics),
            now,
            now,
        ),
    )


def _upsert_validation_task(conn: sqlite3.Connection, task: dict[str, Any]) -> None:
    conn.execute(
        """
        INSERT INTO validation_tasks(
            task_id, source_system, source_report_id, source_report_path, source_paragraph,
            research_topic, symbol, market, signal_to_validate, sample_period, cost_assumption,
            benchmark, status, validation_report_path, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(task_id) DO UPDATE SET
            source_paragraph=excluded.source_paragraph,
            research_topic=excluded.research_topic,
            symbol=excluded.symbol,
            market=excluded.market,
            signal_to_validate=excluded.signal_to_validate,
            sample_period=excluded.sample_period,
            cost_assumption=excluded.cost_assumption,
            benchmark=excluded.benchmark,
            status=CASE
                WHEN validation_tasks.status IN ('已完成', '暂停') THEN validation_tasks.status
                ELSE excluded.status
            END,
            validation_report_path=CASE
                WHEN validation_tasks.validation_report_path <> '' THEN validation_tasks.validation_report_path
                ELSE excluded.validation_report_path
            END,
            updated_at=excluded.updated_at
        """,
        (
            task["task_id"],
            task["source_system"],
            task["source_report_id"],
            task["source_report_path"],
            task["source_paragraph"],
            task["research_topic"],
            task["symbol"],
            task["market"],
            task["signal_to_validate"],
            task["sample_period"],
            task["cost_assumption"],
            task["benchmark"],
            task["status"],
            task["validation_report_path"],
            task["created_at"],
            task["updated_at"],
        ),
    )


def _upsert_pfi_os_result(conn: sqlite3.Connection, result_id: str, row: dict[str, Any], payload: dict[str, Any]) -> None:
    metadata = payload.get("metadata", {}) if isinstance(payload, dict) else {}
    backtest = metadata.get("backtest", {}) if isinstance(metadata, dict) else {}
    quality = payload.get("data_quality", {}) if isinstance(payload, dict) else {}
    cross = payload.get("cross_validation", {}) if isinstance(payload, dict) else {}
    now = _now()
    conn.execute(
        """
        INSERT INTO pfi_os_results(
            result_id, report_path, metadata_path, strategy_id, symbol, market, total_return,
            annualized_return, max_drawdown, sharpe, research_status, decision_quality_score,
            data_quality_status, cross_validation_status, created_at, updated_at, payload_json
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(result_id) DO UPDATE SET
            report_path=excluded.report_path,
            metadata_path=excluded.metadata_path,
            strategy_id=excluded.strategy_id,
            symbol=excluded.symbol,
            market=excluded.market,
            total_return=excluded.total_return,
            annualized_return=excluded.annualized_return,
            max_drawdown=excluded.max_drawdown,
            sharpe=excluded.sharpe,
            research_status=excluded.research_status,
            decision_quality_score=excluded.decision_quality_score,
            data_quality_status=excluded.data_quality_status,
            cross_validation_status=excluded.cross_validation_status,
            updated_at=excluded.updated_at,
            payload_json=excluded.payload_json
        """,
        (
            result_id,
            str(row.get("report_path", "")),
            str(row.get("metadata_path", "")),
            str(row.get("strategy_id", "")),
            str(backtest.get("symbol", "")),
            str(backtest.get("market", "")),
            _float(row.get("total_return")),
            _float(row.get("annualized_return")),
            _float(row.get("max_drawdown")),
            _float(row.get("sharpe")),
            str(row.get("research_status", "")),
            _float(row.get("decision_quality_score")),
            str(quality.get("status", "")) if isinstance(quality, dict) else "",
            str(cross.get("status", "")) if isinstance(cross, dict) else "",
            now,
            now,
            _json_dumps(payload),
        ),
    )


def _upsert_holding(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    source_system = str(row.get("source_system", ""))
    symbol = str(row.get("symbol", ""))
    name = str(row.get("name", ""))
    market = str(row.get("market", ""))
    source_path = str(row.get("source_file", ""))
    updated_at = str(row.get("updated_at", "")) or _now()
    holding_id = _stable_id("holding", source_system, symbol, name, market, source_path)
    conn.execute(
        """
        INSERT INTO holdings_master(
            holding_id, source_system, account, symbol, name, market, asset_type, quantity,
            cost_basis, position_value, unrealized_pnl, weight, as_of, source_path, payload_json, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(holding_id) DO UPDATE SET
            quantity=excluded.quantity,
            cost_basis=excluded.cost_basis,
            position_value=excluded.position_value,
            unrealized_pnl=excluded.unrealized_pnl,
            weight=excluded.weight,
            as_of=excluded.as_of,
            source_path=excluded.source_path,
            payload_json=excluded.payload_json,
            updated_at=excluded.updated_at
        """,
        (
            holding_id,
            source_system,
            str(row.get("account", "DefaultAccount")),
            symbol,
            name,
            market,
            str(row.get("asset_type", "")),
            _float(row.get("quantity")),
            _float(row.get("cost_basis")),
            _float(row.get("position_value")),
            _float(row.get("unrealized_pnl")),
            _float(row.get("weight")),
            str(row.get("updated_at", "")),
            source_path,
            _json_dumps(row),
            updated_at,
        ),
    )


def _upsert_holding_symbol_mapping(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    holding_name = str(row.get("name", "")).strip()
    holding_market = str(row.get("market", "")).strip().upper()
    original_symbol = str(row.get("symbol", "")).strip()
    proxy_symbol = str(row.get("proxy_symbol", "")).strip()
    proxy_market = str(row.get("proxy_market", "")).strip().upper()
    status = str(row.get("status", "")).strip()
    mapping_id = _stable_id("holdingSymbolMapping", holding_name, holding_market, original_symbol, proxy_symbol, proxy_market)
    updated_at = _now()
    payload = {
        "holding_name": holding_name,
        "holding_market": holding_market,
        "original_symbol": original_symbol,
        "proxy_symbol": proxy_symbol,
        "proxy_name": str(row.get("proxy_name", "")).strip(),
        "proxy_market": proxy_market,
        "status": status,
        "confidence": str(row.get("confidence", "")).strip(),
        "reason": str(row.get("reason", "")).strip(),
    }
    conn.execute(
        """
        INSERT INTO holding_symbol_mappings(
            mapping_id, source_system, holding_name, holding_market, original_symbol,
            proxy_symbol, proxy_name, proxy_market, status, confidence, reason, source,
            payload_json, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(mapping_id) DO UPDATE SET
            source_system=excluded.source_system,
            holding_name=excluded.holding_name,
            holding_market=excluded.holding_market,
            original_symbol=excluded.original_symbol,
            proxy_symbol=excluded.proxy_symbol,
            proxy_name=excluded.proxy_name,
            proxy_market=excluded.proxy_market,
            status=excluded.status,
            confidence=excluded.confidence,
            reason=excluded.reason,
            source=excluded.source,
            payload_json=excluded.payload_json,
            updated_at=excluded.updated_at
        """,
        (
            mapping_id,
            "PFIOS",
            holding_name,
            holding_market,
            original_symbol,
            proxy_symbol,
            str(row.get("proxy_name", "")).strip(),
            proxy_market,
            status,
            str(row.get("confidence", "")).strip(),
            str(row.get("reason", "")).strip(),
            "HoldingSymbolMap",
            _json_dumps(payload),
            updated_at,
        ),
    )


def _upsert_portfolio_transaction(conn: sqlite3.Connection, row: dict[str, Any]) -> None:
    source_system = str(row.get("source_system", ""))
    trade_date = str(row.get("trade_date", ""))
    order_time = str(row.get("order_time", ""))
    name = str(row.get("name", ""))
    side = str(row.get("side", ""))
    order_amount = _float(row.get("order_amount"))
    source_path = str(row.get("source_path", ""))
    evidence_frame = str(row.get("evidence_frame", ""))
    transaction_id = str(row.get("transaction_id", "")).strip() or _stable_id(
        "portfolioTxn",
        source_system,
        trade_date,
        order_time,
        name,
        side,
        order_amount,
        source_path,
        evidence_frame,
    )
    updated_at = str(row.get("updated_at", "")) or _now()
    conn.execute(
        """
        INSERT INTO portfolio_transactions(
            transaction_id, source_system, account, trade_date, order_time, timezone, symbol, name,
            market, asset_type, side, order_type, order_amount, confirmed_amount, confirmed_units,
            confirmed_nav, fee, status, quality_status, source_path, evidence_frame, notes,
            payload_json, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(transaction_id) DO UPDATE SET
            source_system=excluded.source_system,
            account=excluded.account,
            trade_date=excluded.trade_date,
            order_time=excluded.order_time,
            timezone=excluded.timezone,
            symbol=excluded.symbol,
            name=excluded.name,
            market=excluded.market,
            asset_type=excluded.asset_type,
            side=excluded.side,
            order_type=excluded.order_type,
            order_amount=excluded.order_amount,
            confirmed_amount=excluded.confirmed_amount,
            confirmed_units=excluded.confirmed_units,
            confirmed_nav=excluded.confirmed_nav,
            fee=excluded.fee,
            status=excluded.status,
            quality_status=excluded.quality_status,
            source_path=excluded.source_path,
            evidence_frame=excluded.evidence_frame,
            notes=excluded.notes,
            payload_json=excluded.payload_json,
            updated_at=excluded.updated_at
        """,
        (
            transaction_id,
            source_system,
            str(row.get("account", "支付宝基金账户")),
            trade_date,
            order_time,
            str(row.get("timezone", "Asia/Shanghai")),
            str(row.get("symbol", "")),
            name,
            str(row.get("market", "CN")),
            str(row.get("asset_type", "fund")),
            side,
            str(row.get("order_type", "")),
            order_amount,
            _float(row.get("confirmed_amount")),
            _float(row.get("confirmed_units")),
            _float(row.get("confirmed_nav")),
            _float(row.get("fee")),
            str(row.get("status", "")),
            str(row.get("quality_status", "")),
            source_path,
            evidence_frame,
            str(row.get("notes", "")),
            _json_dumps(row),
            updated_at,
        ),
    )


def _record_event(
    conn: sqlite3.Connection,
    event_type: str,
    source_system: str,
    target_system: str,
    status: str,
    message: str,
    payload: dict[str, Any] | None = None,
) -> None:
    event_id = _stable_id("event", event_type, source_system, target_system, status, message, _now())
    conn.execute(
        """
        INSERT INTO sync_events(event_id, event_type, source_system, target_system, status, message, payload_json, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (event_id, event_type, source_system, target_system, status, message, _json_dumps(payload or {}), _now()),
    )


def _candidate_paragraphs(text: str) -> list[str]:
    if not text.strip():
        return []
    pieces = [piece.strip() for piece in re.split(r"(?:\n{2,}|[。！？!?]\s+)", text) if piece.strip()]
    candidates = []
    for piece in pieces:
        lowered = piece.lower()
        if len(piece) < 12:
            continue
        if any(keyword in lowered for keyword in VALIDATION_KEYWORDS) or extract_symbols(piece):
            candidates.append(_truncate(piece, 900))
    return candidates


def _research_topic(paragraph: str, report_row: dict[str, Any]) -> str:
    category = str(report_row.get("category", "") or "行研验证")
    clean = _truncate(_compact_whitespace(paragraph), 80)
    return f"{category}：{clean}" if clean else category


def _signal_to_validate(paragraph: str) -> str:
    return _truncate(_compact_whitespace(paragraph), 280)


def _read_docx_text(path: Path) -> str:
    try:
        from docx import Document
    except ModuleNotFoundError:
        return ""
    try:
        document = Document(path)
    except Exception:
        return ""
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _read_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "").strip() for page in reader.pages if page.extract_text())
    except Exception:
        pass
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        return "\n".join((page.extract_text() or "").strip() for page in reader.pages if page.extract_text())
    except Exception:
        pass
    try:
        import fitz

        document = fitz.open(str(path))
        return "\n".join(page.get_text("text").strip() for page in document)
    except Exception:
        pass
    return _read_pdf_ocr_text(path)


def _read_pdf_ocr_text(path: Path) -> str:
    if not _ocr_runtime_available():
        return ""
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except Exception:
        return ""
    texts: list[str] = []
    try:
        document = fitz.open(str(path))
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.frombytes("RGB", [pixmap.width, pixmap.height], pixmap.samples)
            text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
            if text:
                texts.append(text)
    except Exception:
        return ""
    return "\n".join(texts)


def _ocr_runtime_available() -> bool:
    if shutil.which("tesseract") is None:
        return False
    try:
        import fitz  # noqa: F401
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except Exception:
        return False
    return True


def _pdf_parse_warning() -> str:
    if _ocr_runtime_available():
        return "PDF 正文未解析；OCR 引擎可用但本文件没有提取到文字，已仅按文件名和元数据索引。"
    return "PDF 正文未解析；文本层为空且 OCR 引擎未配置，扫描版 PDF 只能按文件名和元数据索引。"


def _default_transaction_roots() -> tuple[Path, ...]:
    configured = os.getenv("PFI_PORTFOLIO_TRANSACTION_DIR", "").strip()
    if configured:
        return tuple(Path(item).expanduser() for item in configured.split(":") if item.strip())
    roots = [DATA_DIR / "holdings" / "imports"]
    ai_alipay = DEFAULT_AI_RESEARCH_ROOT / "data" / "private" / "alipay"
    if ai_alipay.exists():
        roots.append(ai_alipay)
    return tuple(roots)


def _transaction_source_files(path: Path) -> list[Path]:
    expanded = path.expanduser()
    if expanded.is_file():
        return [expanded] if _is_transaction_source_file(expanded) else []
    if not expanded.exists() or not expanded.is_dir():
        return []
    files: list[Path] = []
    for item in sorted(expanded.rglob("*"), key=lambda p: p.stat().st_mtime if p.is_file() else 0, reverse=True):
        if _is_transaction_source_file(item):
            files.append(item)
    return files


def _is_transaction_source_file(path: Path) -> bool:
    if path.is_symlink() or not path.is_file() or path.suffix.lower() != ".csv":
        return False
    name = path.name
    return name in {"trade_ledger.csv", "pending_orders.csv", "orders.csv", "transactions.csv", "confirmed_portfolio_transactions.csv"} or name.startswith("video_trade_candidates_")


def _normalize_transaction_row(row: dict[str, Any], path: Path) -> dict[str, object]:
    side = str(_row_value(row, "side", "方向", "类型", "交易类型")).strip()
    order_type = str(_row_value(row, "order_type", "订单类型")).strip()
    if side == "定投" and not order_type:
        order_type = "recurring"
    elif not order_type:
        order_type = "manual"
    source = str(_row_value(row, "source", "来源")).strip()
    status = str(_row_value(row, "status", "交易状态", "状态")).strip()
    source_path = str(_row_value(row, "source_path", "源文件", "source_file")).strip() or str(path)
    return {
        "transaction_id": str(_row_value(row, "transaction_id", "交易订单号", "id")).strip(),
        "source_system": str(_row_value(row, "source_system")).strip() or _transaction_source_system(path, source),
        "account": str(_row_value(row, "account", "账户")).strip() or "支付宝基金账户",
        "trade_date": str(_row_value(row, "trade_date", "日期", "date")).strip(),
        "order_time": str(_row_value(row, "order_time", "时间", "time")).strip(),
        "timezone": str(_row_value(row, "timezone", "时区")).strip() or "Asia/Shanghai",
        "symbol": str(_row_value(row, "symbol", "代码", "证券代码")).strip(),
        "name": str(_row_value(row, "name", "名称", "基金名称", "标的名称")).strip(),
        "market": str(_row_value(row, "market", "市场")).strip() or "CN",
        "asset_type": str(_row_value(row, "asset_type", "资产类型")).strip() or "fund",
        "side": side,
        "order_type": order_type,
        "order_amount": _float(_row_value(row, "order_amount", "金额", "amount")),
        "confirmed_amount": _float(_row_value(row, "confirmed_amount", "确认金额")),
        "confirmed_units": _float(_row_value(row, "confirmed_units", "确认份额")),
        "confirmed_nav": _float(_row_value(row, "confirmed_nav", "确认净值")),
        "fee": _float(_row_value(row, "fee", "费用", "手续费")),
        "status": status,
        "quality_status": str(_row_value(row, "quality_status", "质量状态")).strip() or _transaction_quality_status(status, source, path),
        "source_path": source_path,
        "evidence_frame": str(_row_value(row, "evidence_frame", "证据帧")).strip(),
        "notes": str(_row_value(row, "notes", "备注")).strip(),
        "updated_at": str(_row_value(row, "updated_at", "更新时间")).strip()
        or datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
    }


def _transaction_source_system(path: Path, source: str) -> str:
    text = f"{path} {source}".lower()
    if "alipay_transaction_csv" in text or "trade_ledger" in text:
        return "支付宝官方交易CSV"
    if "pending_orders" in text:
        return "支付宝待确认订单"
    if "video_trade_candidates" in text or "alipay_video" in text:
        return "支付宝视频候选交易"
    return "量化回测系统导入交易"


def _transaction_quality_status(status: str, source: str, path: Path) -> str:
    text = f"{status} {source} {path}".lower()
    if "不确定" in text or "未清晰" in text or "截断" in text:
        return "CandidateUncertain"
    if "交易成功" in text or "confirmed" in text or "trade_ledger" in text:
        return "Confirmed"
    if "交易进行中" in text or "确认中" in text or "pending" in text:
        return "PendingConfirmation"
    if "video" in text:
        return "VideoCandidate"
    return "Unknown"


def _transaction_quality_priority(status: str) -> int:
    return {
        "Confirmed": 10,
        "PendingConfirmation": 20,
        "VideoCandidate": 30,
        "CandidateUncertain": 40,
        "Unknown": 90,
    }.get(str(status), 90)


def _transaction_source_priority(source_system: str) -> int:
    return {
        "支付宝官方交易CSV": 10,
        "支付宝待确认订单": 20,
        "支付宝视频候选交易": 30,
        "量化回测系统导入交易": 40,
    }.get(str(source_system), 90)


def _transaction_columns() -> list[str]:
    return [
        "transaction_id",
        "source_system",
        "account",
        "trade_date",
        "order_time",
        "timezone",
        "symbol",
        "name",
        "market",
        "asset_type",
        "side",
        "order_type",
        "order_amount",
        "confirmed_amount",
        "confirmed_units",
        "confirmed_nav",
        "fee",
        "status",
        "quality_status",
        "source_path",
        "evidence_frame",
        "notes",
        "updated_at",
    ]


def _row_value(row: dict[str, Any], *keys: str) -> object:
    normalized = {str(key).strip().lower(): value for key, value in row.items()}
    for key in keys:
        if key in row:
            return row[key]
        lower = key.strip().lower()
        if lower in normalized:
            return normalized[lower]
    return ""


def _normalize_symbol(symbol: str) -> str:
    text = symbol.upper().strip()
    if re.fullmatch(r"[036]\d{5}", text):
        suffix = "SH" if text.startswith("6") else "SZ"
        return f"{text}.{suffix}"
    match = re.fullmatch(r"(SH|SZ)\.?(\d{6})", text)
    if match:
        return f"{match.group(2)}.{match.group(1)}"
    match = re.fullmatch(r"(\d{6})\.(SS|SH|SZ)", text)
    if match:
        suffix = "SH" if match.group(2) == "SS" else match.group(2)
        return f"{match.group(1)}.{suffix}"
    match = re.fullmatch(r"HK\.?(\d{4,5})", text)
    if match:
        return f"{match.group(1).zfill(5)}.HK"
    match = re.fullmatch(r"(\d{4,5})\.HK", text)
    if match:
        return f"{match.group(1).zfill(5)}.HK"
    return text


def _summary_text(text: str, fallback: str) -> str:
    compact = _compact_whitespace(text)
    if compact:
        return _truncate(compact, 360)
    return f"未解析正文，仅根据文件名建立索引：{fallback}"


def _hash_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _stable_id(prefix: str, *parts: object) -> str:
    raw = "\n".join(str(part) for part in parts)
    return f"{prefix}_{hashlib.sha1(raw.encode('utf-8')).hexdigest()[:20]}"


def _safe_json_file(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {}
    return payload if isinstance(payload, dict) else {"payload": payload}


def _safe_json_text(value: object) -> dict[str, Any]:
    try:
        payload = json.loads(str(value or "{}"))
    except json.JSONDecodeError:
        return {}
    return payload if isinstance(payload, dict) else {"payload": payload}


def _json_dumps(payload: Any) -> str:
    return json.dumps(payload, ensure_ascii=False, sort_keys=True, default=str)


def _compact_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", str(text or "")).strip()


def _truncate(text: str, limit: int) -> str:
    clean = str(text or "").strip()
    return clean if len(clean) <= limit else f"{clean[: limit - 1]}…"


def _float(value: object) -> float:
    try:
        text = str(value).strip().replace(",", "").replace("元", "").replace("%", "")
        if text == "" or text.lower() == "nan":
            return 0.0
        return float(text)
    except (TypeError, ValueError):
        return 0.0


def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _table_count(conn: sqlite3.Connection, table: str) -> int:
    return int(conn.execute(f"SELECT COUNT(*) FROM {table}").fetchone()[0])


def _validation_queue_key(payload: dict[str, Any]) -> tuple[str, str, str, str]:
    return (
        str(payload.get("source_report", "")),
        str(payload.get("symbol", "")),
        str(payload.get("signal_to_validate", "")),
        str(payload.get("source_paragraph", ""))[:160],
    )


def _normalize_queue_status(status: str) -> str:
    text = str(status or "").strip()
    return text if text in {"待验证", "验证中", "已完成", "暂停"} else "待验证"

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path, PurePosixPath

from docx import Document
from pypdf import PdfReader

from app.services.skill_catalog import extract_skills


ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MAX_EXTRACTED_CHARACTERS = 500_000
MAX_PDF_PAGES = 200
MAX_DOCX_ARCHIVE_ENTRIES = 5_000
MAX_DOCX_UNCOMPRESSED_BYTES = 50 * 1024 * 1024
MAX_DOCX_SINGLE_ENTRY_BYTES = 20 * 1024 * 1024
DATE_PATTERN = re.compile(
    r"\b(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|"
    r"aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?|"
    r"20\d{2}|19\d{2})\b",
    re.IGNORECASE,
)
EMAIL_PATTERN = re.compile(r"[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}", re.IGNORECASE)
PHONE_PATTERN = re.compile(r"(?:\+?\d[\d ()-]{7,}\d)")
URL_PATTERN = re.compile(r"https?://[^\s<>()]+", re.IGNORECASE)


@dataclass
class ParsedResume:
    text: str
    file_type: str
    skills: list[str]
    profile_hints: dict[str, str]
    experiences: list[dict[str, object]]


class ResumeParseError(ValueError):
    pass


def parse_resume(filename: str, data: bytes) -> ParsedResume:
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ResumeParseError("仅支持 PDF、DOCX、TXT 或 Markdown 文件。")
    if not data:
        raise ResumeParseError("文件为空。")

    if suffix == ".pdf":
        text = _parse_pdf(data)
    elif suffix == ".docx":
        text = _parse_docx(data)
    else:
        text = _parse_text(data)

    text = _clean_text(text)
    if len(text) > MAX_EXTRACTED_CHARACTERS:
        raise ResumeParseError("简历文字过多，超过安全处理上限。请上传精简后的求职简历。")
    if len(text) < 80:
        raise ResumeParseError("没有从文件中读取到足够文字。扫描版 PDF 请先转换为可选择文字的 PDF，或改用 DOCX。")

    return ParsedResume(
        text=text,
        file_type=suffix.lstrip("."),
        skills=extract_skills(text),
        profile_hints=_extract_profile_hints(text),
        experiences=_extract_experience_blocks(text),
    )


def _parse_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        if len(reader.pages) > MAX_PDF_PAGES:
            raise ResumeParseError("PDF 页数过多，请上传精简后的求职简历。")
        chunks: list[str] = []
        total = 0
        for page in reader.pages:
            value = page.extract_text() or ""
            total += len(value)
            if total > MAX_EXTRACTED_CHARACTERS:
                raise ResumeParseError("PDF 提取文字过多，请上传精简后的求职简历。")
            chunks.append(value)
        return "\n\n".join(chunks)
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError("PDF 无法读取，请换用 DOCX 或文本版简历。") from exc


def _validate_docx_archive(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            entries = archive.infolist()
            if len(entries) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise ResumeParseError("DOCX 内部文件数量异常，请重新导出简历。")
            total = 0
            for entry in entries:
                normalized = entry.filename.replace("\\", "/")
                path = PurePosixPath(normalized)
                if path.is_absolute() or ".." in path.parts:
                    raise ResumeParseError("DOCX 包含不安全的内部路径，请重新导出简历。")
                if entry.flag_bits & 0x1:
                    raise ResumeParseError("不支持加密的 DOCX，请先解除文件密码。")
                if entry.file_size > MAX_DOCX_SINGLE_ENTRY_BYTES:
                    raise ResumeParseError("DOCX 内部单个对象过大，请移除大型图片或附件。")
                total += entry.file_size
                if total > MAX_DOCX_UNCOMPRESSED_BYTES:
                    raise ResumeParseError("DOCX 解压后体积过大，请移除大型图片或附件。")
    except ResumeParseError:
        raise
    except (OSError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        raise ResumeParseError("DOCX 无法读取，请确认文件没有损坏。") from exc


def _parse_docx(data: bytes) -> str:
    _validate_docx_archive(data)
    try:
        document = Document(io.BytesIO(data))
        chunks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                chunks.append(" | ".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(chunks)
        if len(text) > MAX_EXTRACTED_CHARACTERS:
            raise ResumeParseError("DOCX 提取文字过多，请上传精简后的求职简历。")
        return text
    except ResumeParseError:
        raise
    except Exception as exc:
        raise ResumeParseError("DOCX 无法读取，请确认文件没有损坏。") from exc


def _parse_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise ResumeParseError("文本文件编码无法识别。")


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_profile_hints(text: str) -> dict[str, str]:
    hints: dict[str, str] = {}
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if lines:
        first = lines[0]
        if 1 < len(first.split()) <= 6 and len(first) <= 80 and "@" not in first:
            hints["preferred_name"] = first
    email = EMAIL_PATTERN.search(text)
    if email:
        hints["email"] = email.group(0)
    phone = PHONE_PATTERN.search(text)
    if phone:
        hints["phone"] = re.sub(r"\s+", " ", phone.group(0)).strip()
    for url in URL_PATTERN.findall(text):
        lower = url.lower()
        if "linkedin.com" in lower and "linkedin_url" not in hints:
            hints["linkedin_url"] = url.rstrip(".,;")
        elif "github.com" in lower and "github_url" not in hints:
            hints["github_url"] = url.rstrip(".,;")
        elif "portfolio_url" not in hints:
            hints["portfolio_url"] = url.rstrip(".,;")
    return hints


def _extract_experience_blocks(text: str) -> list[dict[str, object]]:
    section_aliases = {
        "experience": ("experience", "employment", "work history", "professional experience", "实习", "工作经历"),
        "project": ("projects", "project experience", "selected projects", "项目经历", "项目"),
    }
    lines = [line.strip() for line in text.splitlines()]
    sections: list[tuple[str, int, int]] = []
    heading_positions: list[tuple[str, int]] = []

    for index, line in enumerate(lines):
        normalized = re.sub(r"[^a-z\u4e00-\u9fff ]", "", line.lower()).strip()
        for category, aliases in section_aliases.items():
            if normalized in aliases or any(normalized == alias for alias in aliases):
                heading_positions.append((category, index))
                break

    for pos, (category, start) in enumerate(heading_positions):
        end = heading_positions[pos + 1][1] if pos + 1 < len(heading_positions) else len(lines)
        sections.append((category, start + 1, end))

    candidates: list[dict[str, object]] = []
    for category, start, end in sections:
        raw = lines[start:end]
        blocks = _group_blocks(raw)
        for block in blocks:
            candidate = _block_to_experience(category, block)
            if candidate:
                candidates.append(candidate)
            if len(candidates) >= 10:
                return candidates
    return candidates


def _group_blocks(lines: list[str]) -> list[list[str]]:
    blocks: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        if not line:
            if current:
                blocks.append(current)
                current = []
            continue
        if current and DATE_PATTERN.search(line) and len(current) >= 3:
            blocks.append(current)
            current = [line]
        else:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def _block_to_experience(category: str, block: list[str]) -> dict[str, object] | None:
    joined = "\n".join(block).strip()
    if len(joined) < 70:
        return None
    title = block[0][:220]
    organization = block[1][:220] if len(block) > 1 and len(block[1]) < 160 else ""
    date_line = next((line for line in block[:4] if DATE_PATTERN.search(line)), "")
    description_lines = block[2:] if organization else block[1:]
    description = "\n".join(description_lines).strip()[:5000]
    return {
        "category": category,
        "title": title,
        "organization": organization,
        "date_range": date_line[:120],
        "description": description,
        "tags": extract_skills(joined),
    }

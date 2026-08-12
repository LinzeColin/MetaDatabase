from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def _text(value: Any) -> str:
    return str(value or "").strip()


def build_tailored_resume_docx(content: dict[str, Any], parsed: dict[str, Any]) -> bytes:
    """Build a job-specific DOCX from confirmed resume facts only.

    The document intentionally does not reconstruct a visually identical copy of
    the source resume. It produces a clean ATS-oriented draft and labels all gaps
    so the user can review it before upload.
    """
    materials = content.get("materials") or {}
    tailored = materials.get("tailored_cv") or {}
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    candidate_name = _text(tailored.get("candidate_name") or parsed.get("candidate_name"))
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(candidate_name or "岗位定制简历")
    run.bold = True
    run.font.size = Pt(18)

    target = document.add_paragraph()
    target.alignment = WD_ALIGN_PARAGRAPH.CENTER
    target.add_run(f"目标岗位：{_text(content.get('job_title'))}｜{_text(content.get('company'))}").bold = True

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("本文件由已上传简历中的事实派生；提交前必须由本人核对。")
    note_run.italic = True
    note_run.font.size = Pt(8.5)

    summary = _text(tailored.get("summary"))
    if summary:
        document.add_heading("岗位适配摘要", level=1)
        document.add_paragraph(summary)

    skills = [_text(item) for item in tailored.get("skills", []) if _text(item)]
    if skills:
        document.add_heading("核心技能", level=1)
        document.add_paragraph(" · ".join(skills[:24]))

    bullets = [_text(item) for item in tailored.get("bullets", []) if _text(item)]
    document.add_heading("相关经历", level=1)
    if bullets:
        for item in bullets:
            document.add_paragraph(item, style="List Bullet")
    else:
        document.add_paragraph("当前所选简历没有可直接复述的相关经历，请先补充真实内容。")

    education = [_text(item) for item in tailored.get("education", []) if _text(item)]
    if education:
        document.add_heading("教育背景", level=1)
        for item in education:
            document.add_paragraph(item, style="List Bullet")

    credentials = [_text(item) for item in tailored.get("credentials", []) if _text(item)]
    if credentials:
        document.add_heading("专业资质", level=1)
        document.add_paragraph(" · ".join(credentials))

    matched = [_text(item) for item in tailored.get("matched_terms", []) if _text(item)]
    if matched:
        document.add_heading("已覆盖的岗位要求", level=1)
        document.add_paragraph(" · ".join(matched))

    gaps = [_text(item) for item in tailored.get("gaps", []) if _text(item)]
    if gaps:
        document.add_heading("提交前待确认", level=1)
        for item in gaps:
            document.add_paragraph(item, style="List Bullet")

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("JobHuntBot｜事实约束的岗位定制草稿").font.size = Pt(8)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()

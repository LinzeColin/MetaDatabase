from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader

from .career_intelligence import (
    SKILL_TERMS,
    ROLE_RULES,
    detect_credentials,
    detect_role_family,
    detect_skills,
    estimate_experience_years,
    infer_professional_facts,
)

LOCATION_TERMS = [
    "Sydney", "Melbourne", "Brisbane", "Perth", "Canberra", "Adelaide",
    "Remote Australia", "Australia", "悉尼", "墨尔本", "布里斯班", "珀斯", "堪培拉", "远程",
]

EXPERIENCE_HINTS = (
    "intern", "analyst", "accountant", "auditor", "finance", "investment", "treasury", "tax",
    "lawyer", "solicitor", "legal", "paralegal", "clerk", "counsel", "compliance", "contract",
    "manager", "assistant", "project", "research", "operations", "consultant", "associate",
    "实习", "分析", "会计", "审计", "金融", "法律", "律师", "法务", "合规", "合同", "项目", "运营",
)

EDUCATION_HINTS = (
    "university", "bachelor", "master", "degree", "unsw", "college", "juris doctor", "llb", "jd",
    "大学", "本科", "硕士", "学士", "博士", "法学",
)


class ResumeError(ValueError):
    pass


def _decode_plain_text(data: bytes) -> str:
    """Decode normal UTF text and common Chinese resume encodings safely."""
    if data.startswith((b"\xff\xfe", b"\xfe\xff")):
        return data.decode("utf-16")
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_text(filename: str, content_type: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" or content_type == "application/pdf":
        reader = PdfReader(io.BytesIO(data))
        text = "\n".join((page.extract_text() or "") for page in reader.pages)
    elif suffix == ".docx" or "wordprocessingml" in content_type:
        doc = Document(io.BytesIO(data))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
    elif suffix in {".txt", ".md"} or content_type.startswith("text/"):
        text = _decode_plain_text(data)
    else:
        raise ResumeError("支持 PDF、DOCX、TXT 和 Markdown 简历。")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    if len(text.strip()) < 80:
        raise ResumeError("没有从文件中读取到足够的简历文字，请换用可选择文字的 PDF、DOCX 或 TXT。")
    return text.strip()


def _sentences(text: str) -> list[str]:
    items: list[str] = []
    for raw in text.splitlines():
        line = raw.strip(" \t•▪●-–—")
        if 12 <= len(line) <= 700:
            items.append(line)
    return items


def _role_scores(text: str) -> dict[str, int]:
    lower = text.casefold()
    scores: dict[str, int] = {}
    for role, terms in ROLE_RULES.items():
        score = sum(lower.count(term.casefold()) for term in terms)
        scores[role] = score
    detected = detect_role_family(text)
    if detected != "Other":
        scores[detected] = scores.get(detected, 0) + 4
    return scores


def _extract_name(text: str) -> str:
    # Keep this intentionally conservative. The name is only used inside the
    # user's generated document and never becomes a job-match fact.
    for line in text.splitlines()[:6]:
        cleaned = line.strip()
        if 2 <= len(cleaned) <= 70 and not re.search(r"[@|/\\]|resume|curriculum|简历", cleaned, re.I):
            if len(cleaned.split()) <= 6:
                return cleaned
    return ""


def parse_resume(text: str) -> dict[str, Any]:
    lower = text.casefold()
    skills = sorted(set(detect_skills(text)))
    role_scores = _role_scores(text)
    role_families = [
        role for role, score in sorted(role_scores.items(), key=lambda item: (-item[1], item[0]))
        if score > 0
    ][:5]
    education: list[str] = []
    experiences: list[str] = []
    for line in _sentences(text):
        l = line.casefold()
        if any(term in l for term in EDUCATION_HINTS):
            education.append(line)
        if any(term in l for term in EXPERIENCE_HINTS):
            experiences.append(line)

    credentials = detect_credentials(text)
    professional = infer_professional_facts(text)
    locations = []
    for loc in LOCATION_TERMS:
        if loc.casefold() in lower:
            canonical = {
                "悉尼": "Sydney", "墨尔本": "Melbourne", "布里斯班": "Brisbane", "珀斯": "Perth",
                "堪培拉": "Canberra", "远程": "Remote Australia",
            }.get(loc, loc)
            if canonical not in locations:
                locations.append(canonical)

    # Keep keywords bounded and deterministic. This is a discovery seed, not a
    # claim that every token is a skill.
    raw_keywords = [
        token for token in re.findall(r"[A-Za-z][A-Za-z0-9+#.\-]{2,}", text)
        if len(token) < 30
    ]
    keywords = sorted(set(skills + role_families + credentials + raw_keywords), key=str.casefold)[:100]
    years = professional.get("experience_years")
    summary_parts = [
        f"识别到 {len(skills)} 项技能",
        f"{len(experiences[:12])} 段可能相关经历",
        f"{len(role_families)} 个候选岗位方向",
    ]
    if years is not None:
        summary_parts.append(f"约 {years:g} 年相关经验（待本人确认）")
    if credentials:
        summary_parts.append("专业资质：" + "、".join(credentials[:6]))

    return {
        "candidate_name": _extract_name(text),
        "skills": skills,
        "role_families": role_families,
        "education": education[:10],
        "experiences": experiences[:16],
        "locations": locations,
        "keywords": keywords,
        "experience_years": years,
        "professional_credentials": credentials,
        "legal_admission": professional.get("legal_admission", "uncertain"),
        "practising_certificate": professional.get("practising_certificate", "uncertain"),
        "education_levels": professional.get("education_levels", []),
        "summary": "系统从简历中" + "、".join(summary_parts) + "。",
    }


def profile_draft(parsed: dict[str, Any]) -> dict[str, Any]:
    # Do not invent locations or work modes. The user confirms them on the next
    # page. Empty is safer than a polished but false profile.
    return {
        "primary_role_families": parsed.get("role_families", [])[:3],
        "secondary_role_families": parsed.get("role_families", [])[3:],
        "target_locations": parsed.get("locations", []),
        "work_mode": [],
        "skills": parsed.get("skills", []),
        "keywords": parsed.get("keywords", [])[:50],
        "experience_years": parsed.get("experience_years"),
        "professional_credentials": parsed.get("professional_credentials", []),
        "credentials_confirmed": False,
        "education_levels": parsed.get("education_levels", []),
        "legal_admission": parsed.get("legal_admission", "uncertain"),
        "practising_certificate": parsed.get("practising_certificate", "uncertain"),
        "work_authorization": "",
        "sponsorship_now": "",
        "sponsorship_future": "",
        "relocation": "",
        "available_start": "",
        "avoid_roles": [],
        "avoid_industries": [],
    }


def experience_records(parsed: dict[str, Any]) -> list[dict[str, str]]:
    return [
        {"title": f"经历 {idx}", "detail": detail, "kind": "experience", "strength": "medium"}
        for idx, detail in enumerate(parsed.get("experiences", [])[:12], start=1)
    ]
